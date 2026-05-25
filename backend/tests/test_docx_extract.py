from __future__ import annotations

from pathlib import Path

import pytest


def _build_minimal_docx(
    path: Path,
    *,
    paragraphs: tuple[tuple[str, str | None], ...] = (
        ("Jane Doe", "Title"),
        ("Senior Engineer", None),
        ("Experience", "Heading 1"),
        ("Acme Corp — Staff Engineer", "Heading 2"),
        ("Built distributed systems.", "List Bullet"),
        ("Shipped revenue-generating features.", "List Bullet"),
    ),
    table_rows: tuple[tuple[str, ...], ...] | None = None,
) -> None:
    """Generate a minimal real .docx fixture via python-docx.

    Keeping the helper inside the test module avoids committing binary
    fixtures; the resulting file is a real Word package the extractor
    can open with ``python-docx``.
    """
    from docx import Document

    doc = Document()
    for text, style in paragraphs:
        if style is not None:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
    if table_rows:
        cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for r_idx, row in enumerate(table_rows):
            for c_idx, value in enumerate(row):
                table.cell(r_idx, c_idx).text = value
    doc.save(str(path))


def test_find_source_resume_docx_matches_accepted_names(tmp_path: Path):
    from app.docx_extract import (
        ACCEPTED_RESUME_DOCX_NAMES,
        find_source_resume_docx,
    )

    assert find_source_resume_docx(tmp_path) is None
    (tmp_path / "unrelated.docx").write_bytes(b"x")
    assert find_source_resume_docx(tmp_path) is None

    # First match wins, in declared order.
    expected = tmp_path / ACCEPTED_RESUME_DOCX_NAMES[0]
    (tmp_path / ACCEPTED_RESUME_DOCX_NAMES[1]).write_bytes(b"y")
    expected.write_bytes(b"z")
    assert find_source_resume_docx(tmp_path) == expected


def test_extract_master_resume_writes_extracted_markdown(tmp_path: Path):
    from app.docx_extract import (
        EXTRACTED_FILENAME,
        extract_master_resume_if_present,
    )

    _build_minimal_docx(tmp_path / "master_resume.docx")
    result = extract_master_resume_if_present(tmp_path)

    assert result.docx_found is True
    assert result.extracted is True
    assert result.error_message is None

    extracted = (tmp_path / EXTRACTED_FILENAME).read_text(encoding="utf-8")
    # The extracted markdown must point back at the source DOCX filename
    # (relative to the run input directory) so a reader can trace claims
    # back to the formatting source.
    assert "Source DOCX: input/master_resume.docx" in extracted
    assert "# Extracted Master Resume" in extracted
    assert "## Extracted Text" in extracted
    # Headings + bullets preserved in a recognizable form.
    assert "# Jane Doe" in extracted  # Title → level 1
    assert "# Experience" in extracted
    assert "## Acme Corp" in extracted
    assert "- Built distributed systems." in extracted


def test_extract_master_resume_preserves_table_rows(tmp_path: Path):
    from app.docx_extract import (
        EXTRACTED_FILENAME,
        extract_master_resume_if_present,
    )

    _build_minimal_docx(
        tmp_path / "resume.docx",
        paragraphs=(("Skills", "Heading 1"),),
        table_rows=(
            ("Language", "Years"),
            ("Python", "8"),
            ("Rust", "2"),
        ),
    )
    result = extract_master_resume_if_present(tmp_path)
    assert result.extracted is True

    extracted = (tmp_path / EXTRACTED_FILENAME).read_text(encoding="utf-8")
    assert "| Language | Years |" in extracted
    assert "| --- | --- |" in extracted
    assert "| Python | 8 |" in extracted


def test_extract_master_resume_no_docx_returns_empty_result(tmp_path: Path):
    from app.docx_extract import (
        EXTRACTED_FILENAME,
        EXTRACTION_ERROR_FILENAME,
        extract_master_resume_if_present,
    )

    (tmp_path / "master_resume.md").write_text("# md only\n", encoding="utf-8")
    result = extract_master_resume_if_present(tmp_path)

    assert result.docx_found is False
    assert result.extracted is False
    assert result.failed is False
    assert result.has_other_resume_source is True
    # The extractor must not write anything when no DOCX is present.
    assert not (tmp_path / EXTRACTED_FILENAME).exists()
    assert not (tmp_path / EXTRACTION_ERROR_FILENAME).exists()


def test_extract_master_resume_invalid_docx_writes_error_file(tmp_path: Path):
    from app.docx_extract import (
        EXTRACTED_FILENAME,
        EXTRACTION_ERROR_FILENAME,
        extract_master_resume_if_present,
    )

    # Bytes that are clearly not a valid ZIP / Word package — python-docx
    # will refuse to open this.
    (tmp_path / "master_resume.docx").write_bytes(b"not a real docx file")
    (tmp_path / "master_resume.md").write_text("# md fallback\n", encoding="utf-8")

    result = extract_master_resume_if_present(tmp_path)
    assert result.docx_found is True
    assert result.extracted is False
    assert result.failed is True
    assert result.has_other_resume_source is True
    assert result.error_message  # non-empty

    error_text = (tmp_path / EXTRACTION_ERROR_FILENAME).read_text(encoding="utf-8")
    assert "Source DOCX: input/master_resume.docx" in error_text
    assert "Reason:" in error_text
    # No stale extracted file should appear.
    assert not (tmp_path / EXTRACTED_FILENAME).exists()


def test_extract_master_resume_failure_without_markdown_source(tmp_path: Path):
    """Extraction failure must be visible to the caller even when no
    fallback markdown resume exists; the worker uses this signal to
    fail the run loudly rather than tailoring blind."""
    from app.docx_extract import extract_master_resume_if_present

    (tmp_path / "master_resume.docx").write_bytes(b"corrupted")
    result = extract_master_resume_if_present(tmp_path)

    assert result.failed is True
    assert result.has_other_resume_source is False


def test_extract_master_resume_clears_stale_error_on_success(tmp_path: Path):
    from app.docx_extract import (
        EXTRACTION_ERROR_FILENAME,
        extract_master_resume_if_present,
    )

    # Prior failure left an error file behind.
    (tmp_path / EXTRACTION_ERROR_FILENAME).write_text(
        "old error\n", encoding="utf-8"
    )
    _build_minimal_docx(tmp_path / "master_resume.docx")

    result = extract_master_resume_if_present(tmp_path)
    assert result.extracted is True
    # A successful extraction must remove the stale error file so the
    # operator does not see a misleading diagnostic from a previous run.
    assert not (tmp_path / EXTRACTION_ERROR_FILENAME).exists()


def test_extract_master_resume_clears_stale_extracted_on_failure(tmp_path: Path):
    from app.docx_extract import (
        EXTRACTED_FILENAME,
        extract_master_resume_if_present,
    )

    (tmp_path / EXTRACTED_FILENAME).write_text(
        "stale extracted content\n", encoding="utf-8"
    )
    (tmp_path / "master_resume.docx").write_bytes(b"garbage")

    result = extract_master_resume_if_present(tmp_path)
    assert result.failed is True
    # The previous extracted file must not be left behind to look authoritative.
    assert not (tmp_path / EXTRACTED_FILENAME).exists()
