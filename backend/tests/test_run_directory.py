from __future__ import annotations

import json
import shutil
from datetime import datetime, timezone
from pathlib import Path

import pytest


CANDIDATE_FILES = (
    "candidate_profile.md",
    "skills_inventory.md",
    "tailoring_preferences.md",
    "resume_dos_and_donts.md",
)

EXPECTED_INPUT_FILES = {
    "job_description.md",
    "master_resume.md",
    "evidence_bank.md",
    "evidence_sources_index.md",
    "candidate_profile.md",
    "project_notes.md",
    "skills_inventory.md",
    "tailoring_preferences.md",
    "resume_dos_and_donts.md",
    "tailoring_prompt.md",
}


@pytest.fixture()
def fixture_layout(tmp_path: Path) -> dict[str, Path]:
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\nstable content\n", encoding="utf-8")
    notes_dir = candidate_root / "project_notes"
    notes_dir.mkdir()
    (notes_dir / "alpha.md").write_text("alpha note\n", encoding="utf-8")
    (notes_dir / "beta.md").write_text("beta note\n", encoding="utf-8")

    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text(
        "# Prompt\nRead inputs and write outputs.\n", encoding="utf-8"
    )

    runs_root = tmp_path / "runs"
    runs_root.mkdir()

    return {
        "candidate_root": candidate_root,
        "prompts_root": prompts_root,
        "runs_root": runs_root,
    }


def _make_objects():
    from app.models import EvidenceBank, Job, MasterResume

    job = Job(
        id="job-1",
        source_platform="linkedin",
        external_url="https://example.com/job/1",
        company="Acme",
        title="ML Engineer",
        location="Remote",
        description_text="Build ML systems.",
        application_method="form",
    )
    resume = MasterResume(
        id="resume-1",
        name="main",
        content_markdown="# Master Resume\nExperience: ...\n",
    )
    evidence = EvidenceBank(
        id="evidence-1",
        name="main",
        content_markdown="# Evidence\nProject X: shipped.\n",
    )
    return job, resume, evidence


def test_create_run_directory_layout(fixture_layout):
    from app.run_directory import create_run_directory

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )

    assert info.run_dir.is_dir()
    assert (info.run_dir / "input").is_dir()
    assert (info.run_dir / "output").is_dir()
    assert (info.run_dir / "metadata.json").is_file()

    present = {p.name for p in (info.run_dir / "input").iterdir() if p.is_file()}
    assert present == EXPECTED_INPUT_FILES


def test_project_notes_concatenated_from_directory(fixture_layout):
    from app.run_directory import create_run_directory

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )
    notes = (info.run_dir / "input" / "project_notes.md").read_text(encoding="utf-8")
    assert "alpha note" in notes
    assert "beta note" in notes


def test_hashes_stable_across_runs(fixture_layout):
    from app.run_directory import create_run_directory

    job, resume, evidence = _make_objects()
    fixed_now = datetime(2026, 1, 1, tzinfo=timezone.utc)

    a = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
        now=fixed_now,
    )
    b = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
        now=fixed_now,
    )
    assert a.run_id != b.run_id
    assert a.prompt_hash == b.prompt_hash
    assert a.input_hash == b.input_hash


def test_metadata_json_well_formed(fixture_layout):
    from app.run_directory import EXPECTED_OUTPUTS, create_run_directory

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )
    metadata = json.loads((info.run_dir / "metadata.json").read_text(encoding="utf-8"))

    required = {
        "run_id",
        "job_id",
        "master_resume_id",
        "capture_method",
        "created_at",
        "input_files",
        "expected_outputs",
        "prompt_hash",
        "input_hash",
    }
    assert required <= set(metadata.keys())
    assert metadata["run_id"] == info.run_id
    assert metadata["job_id"] == "job-1"
    assert metadata["master_resume_id"] == "resume-1"
    assert metadata["evidence_bank_id"] == "evidence-1"
    assert metadata["expected_outputs"] == list(EXPECTED_OUTPUTS)
    assert metadata["prompt_hash"] == info.prompt_hash
    assert metadata["input_hash"] == info.input_hash
    # Every input file should appear in the input_files map with a sha256 hex.
    assert set(metadata["input_files"].keys()) == EXPECTED_INPUT_FILES
    for digest in metadata["input_files"].values():
        assert len(digest) == 64
        int(digest, 16)  # parseable as hex


def test_evidence_bank_optional(fixture_layout):
    from app.run_directory import create_run_directory

    job, resume, _ = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=None,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )
    text = (info.run_dir / "input" / "evidence_bank.md").read_text(encoding="utf-8")
    assert "(none provided)" in text
    metadata = json.loads((info.run_dir / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["evidence_bank_id"] is None


def test_missing_candidate_context_errors(fixture_layout, tmp_path):
    from app.run_directory import RunDirectoryError, create_run_directory

    job, resume, _ = _make_objects()
    bare_candidate_root = tmp_path / "empty_candidate"
    bare_candidate_root.mkdir()  # exists but missing required files

    with pytest.raises(RunDirectoryError, match="candidate_profile.md"):
        create_run_directory(
            job=job,
            master_resume=resume,
            evidence_bank=None,
            candidate_context_root=bare_candidate_root,
            runs_root=fixture_layout["runs_root"],
            runtime_prompts_root=fixture_layout["prompts_root"],
        )


def test_new_run_defaults_to_auto_tailoring_method(fixture_layout):
    from app.run_directory import (
        DEFAULT_RUN_STATUS,
        DEFAULT_TAILORING_METHOD,
        create_run_directory,
        get_run_status,
        get_tailoring_method,
    )

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )

    metadata = json.loads((info.run_dir / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["tailoring_method"] == "auto"
    assert metadata["status"] == "created"
    assert metadata["updated_at"] == metadata["created_at"]
    assert DEFAULT_TAILORING_METHOD == "auto"
    assert DEFAULT_RUN_STATUS == "created"

    assert get_tailoring_method(info.run_dir) == "auto"
    assert get_run_status(info.run_dir) == "created"


def test_create_run_directory_accepts_word_handoff_method(fixture_layout):
    from app.run_directory import create_run_directory, get_tailoring_method

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
        tailoring_method="word_handoff",
    )

    assert get_tailoring_method(info.run_dir) == "word_handoff"
    metadata = json.loads((info.run_dir / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["tailoring_method"] == "word_handoff"


def test_create_run_directory_rejects_invalid_method(fixture_layout):
    from app.run_directory import RunDirectoryError, create_run_directory

    job, resume, evidence = _make_objects()
    with pytest.raises(RunDirectoryError, match="invalid tailoring_method"):
        create_run_directory(
            job=job,
            master_resume=resume,
            evidence_bank=evidence,
            candidate_context_root=fixture_layout["candidate_root"],
            runs_root=fixture_layout["runs_root"],
            runtime_prompts_root=fixture_layout["prompts_root"],
            tailoring_method="bogus",
        )


def test_set_tailoring_method_roundtrip_and_rejects_invalid(fixture_layout):
    from app.run_directory import (
        RunDirectoryError,
        create_run_directory,
        get_tailoring_method,
        set_tailoring_method,
    )

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )

    set_tailoring_method(info.run_dir, "word_handoff")
    assert get_tailoring_method(info.run_dir) == "word_handoff"

    set_tailoring_method(info.run_dir, "auto")
    assert get_tailoring_method(info.run_dir) == "auto"

    with pytest.raises(RunDirectoryError, match="invalid tailoring_method"):
        set_tailoring_method(info.run_dir, "telegram")


def test_set_run_status_roundtrips_all_allowed_statuses(fixture_layout):
    from app.run_directory import (
        ALLOWED_RUN_STATUSES,
        RunDirectoryError,
        create_run_directory,
        get_run_status,
        set_run_status,
    )

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )

    expected = {
        "created",
        "input_ready",
        "auto_tailoring_running",
        "auto_tailoring_failed",
        "auto_tailoring_complete",
        "word_handoff_ready",
        "waiting_for_word_result",
        "word_result_imported",
        "validation_failed",
        "completed",
        "failed",
    }
    assert expected <= set(ALLOWED_RUN_STATUSES)

    for status in expected:
        set_run_status(info.run_dir, status)
        assert get_run_status(info.run_dir) == status

    with pytest.raises(RunDirectoryError, match="invalid run status"):
        set_run_status(info.run_dir, "shipped_to_mars")


def test_status_updates_bump_updated_at(fixture_layout):
    from datetime import timedelta

    from app.run_directory import create_run_directory, set_run_status

    job, resume, evidence = _make_objects()
    fixed_now = datetime(2026, 1, 1, tzinfo=timezone.utc)
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
        now=fixed_now,
    )

    set_run_status(info.run_dir, "input_ready", now=fixed_now + timedelta(seconds=30))
    metadata = json.loads((info.run_dir / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["updated_at"] != metadata["created_at"]
    assert metadata["updated_at"].startswith("2026-01-01T00:00:30")


def test_create_run_directory_stamps_default_llm_provider(fixture_layout):
    from app.run_directory import (
        DEFAULT_LLM_PROVIDER,
        create_run_directory,
        get_llm_provider,
    )

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )

    metadata = json.loads((info.run_dir / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["llm_provider"] == "claude_code"
    assert DEFAULT_LLM_PROVIDER == "claude_code"
    assert get_llm_provider(info.run_dir) == "claude_code"


def test_create_run_directory_accepts_explicit_llm_provider(fixture_layout):
    from app.run_directory import create_run_directory, get_llm_provider

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
        llm_provider="codex",
    )

    metadata = json.loads((info.run_dir / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["llm_provider"] == "codex"
    assert get_llm_provider(info.run_dir) == "codex"


def test_create_run_directory_word_handoff_uses_sentinel(fixture_layout):
    """Word-handoff runs do not invoke a backend CLI; the field must still be
    present so metadata stays self-describing, and it must carry the
    ``claude_for_word`` sentinel regardless of any caller-supplied value."""
    from app.run_directory import create_run_directory, get_llm_provider

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
        tailoring_method="word_handoff",
        llm_provider="codex",  # intentionally lying — must be overridden
    )

    metadata = json.loads((info.run_dir / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["llm_provider"] == "claude_for_word"
    assert get_llm_provider(info.run_dir) == "claude_for_word"


def test_create_run_directory_rejects_unknown_llm_provider(fixture_layout):
    from app.run_directory import RunDirectoryError, create_run_directory

    job, resume, evidence = _make_objects()
    with pytest.raises(RunDirectoryError, match="invalid llm_provider"):
        create_run_directory(
            job=job,
            master_resume=resume,
            evidence_bank=evidence,
            candidate_context_root=fixture_layout["candidate_root"],
            runs_root=fixture_layout["runs_root"],
            runtime_prompts_root=fixture_layout["prompts_root"],
            llm_provider="bogus",
        )


def test_legacy_metadata_without_llm_provider_reads_as_claude_code(fixture_layout):
    """Metadata predating ADR-009 must read back as the claude_code default."""
    from app.run_directory import (
        DEFAULT_LLM_PROVIDER,
        create_run_directory,
        get_llm_provider,
    )

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )

    metadata_path = info.run_dir / "metadata.json"
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    metadata.pop("llm_provider", None)
    metadata_path.write_text(
        json.dumps(metadata, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )

    assert get_llm_provider(info.run_dir) == DEFAULT_LLM_PROVIDER


def test_legacy_metadata_without_tailoring_method_reads_as_auto(fixture_layout):
    """Metadata predating this field must read back as the auto default."""
    from app.run_directory import (
        DEFAULT_RUN_STATUS,
        DEFAULT_TAILORING_METHOD,
        create_run_directory,
        get_run_status,
        get_tailoring_method,
    )

    job, resume, evidence = _make_objects()
    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
    )

    # Rewrite metadata.json without the new fields to mimic a legacy run.
    metadata_path = info.run_dir / "metadata.json"
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    metadata.pop("tailoring_method", None)
    metadata.pop("status", None)
    metadata.pop("updated_at", None)
    metadata_path.write_text(
        json.dumps(metadata, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )

    assert get_tailoring_method(info.run_dir) == DEFAULT_TAILORING_METHOD
    assert get_run_status(info.run_dir) == DEFAULT_RUN_STATUS


def test_post_run_missing_master_resume_returns_404(client, tmp_path, monkeypatch):
    # Prime a candidate context + prompts dir + runs dir so the router could
    # succeed if the resume existed.
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\n", encoding="utf-8")
    (candidate_root / "project_notes.md").write_text("notes\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text("# prompt\n", encoding="utf-8")
    runs_root = tmp_path / "runs"
    runs_root.mkdir()

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build things",
        },
    ).json()

    resp = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": "does-not-exist"},
    )
    assert resp.status_code == 404
    assert "master resume" in resp.json()["detail"].lower()


def _make_docx_file(path: Path, body: str = "Real DOCX body content.") -> None:
    """Helper that writes a real ``.docx`` via python-docx."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Header", style="Heading 1")
    doc.add_paragraph(body)
    doc.save(str(path))


def test_create_run_directory_copies_docx_when_path_provided(fixture_layout, tmp_path):
    """Passing ``master_resume_docx_path`` must stage the file in ``input/``."""
    from app.run_directory import create_run_directory

    job, resume, evidence = _make_objects()
    docx_src = tmp_path / "source.docx"
    _make_docx_file(docx_src)

    info = create_run_directory(
        job=job,
        master_resume=resume,
        evidence_bank=evidence,
        candidate_context_root=fixture_layout["candidate_root"],
        runs_root=fixture_layout["runs_root"],
        runtime_prompts_root=fixture_layout["prompts_root"],
        master_resume_docx_path=docx_src,
    )

    staged = info.run_dir / "input" / "master_resume.docx"
    assert staged.is_file()
    assert staged.read_bytes() == docx_src.read_bytes()
    # The markdown sibling is still written so the input set stays uniform.
    assert (info.run_dir / "input" / "master_resume.md").is_file()


def test_create_run_directory_rejects_missing_docx_path(fixture_layout, tmp_path):
    from app.run_directory import RunDirectoryError, create_run_directory

    job, resume, evidence = _make_objects()
    with pytest.raises(RunDirectoryError, match="master_resume_docx_path"):
        create_run_directory(
            job=job,
            master_resume=resume,
            evidence_bank=evidence,
            candidate_context_root=fixture_layout["candidate_root"],
            runs_root=fixture_layout["runs_root"],
            runtime_prompts_root=fixture_layout["prompts_root"],
            master_resume_docx_path=tmp_path / "missing.docx",
        )


def test_post_run_with_filesystem_docx_stages_docx_in_input(
    client, tmp_path, monkeypatch
):
    """End-to-end: creating a run with a discovered DOCX must copy it
    into ``runs/<id>/input/master_resume.docx`` and keep the existing
    markdown contract intact."""
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\n", encoding="utf-8")
    (candidate_root / "project_notes.md").write_text("notes\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text("# prompt\n", encoding="utf-8")
    runs_root = tmp_path / "runs"
    runs_root.mkdir()
    master_resumes_root = candidate_root / "master_resumes"
    master_resumes_root.mkdir()
    docx_path = master_resumes_root / "calvin.docx"
    _make_docx_file(docx_path, body="ROUND_TRIP_DOCX_BODY")

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))
    monkeypatch.setenv("JOBAPPLY_MASTER_RESUMES_ROOT", str(master_resumes_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build things",
        },
    ).json()

    listed = client.get("/master-resumes").json()
    fs_entry = next(e for e in listed if e["source"] == "filesystem")

    resp = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": fs_entry["id"]},
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["master_resume_id"] == fs_entry["id"]
    run_dir = Path(body["run_dir"])
    staged_docx = run_dir / "input" / "master_resume.docx"
    assert staged_docx.is_file()
    assert staged_docx.read_bytes() == docx_path.read_bytes()
    # The markdown sibling is also written with the extracted body.
    md_text = (run_dir / "input" / "master_resume.md").read_text(encoding="utf-8")
    assert "ROUND_TRIP_DOCX_BODY" in md_text

    shutil.rmtree(run_dir, ignore_errors=True)


def test_post_run_with_filesystem_md_writes_master_resume_md(
    client, tmp_path, monkeypatch
):
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\n", encoding="utf-8")
    (candidate_root / "project_notes.md").write_text("notes\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text("# prompt\n", encoding="utf-8")
    runs_root = tmp_path / "runs"
    runs_root.mkdir()
    master_resumes_root = candidate_root / "master_resumes"
    master_resumes_root.mkdir()
    (master_resumes_root / "industry_ml.md").write_text(
        "# Industry ML Resume\nSENTINEL_MD\n", encoding="utf-8"
    )

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))
    monkeypatch.setenv("JOBAPPLY_MASTER_RESUMES_ROOT", str(master_resumes_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build",
        },
    ).json()
    listed = client.get("/master-resumes").json()
    fs_entry = next(e for e in listed if e["source"] == "filesystem")

    resp = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": fs_entry["id"]},
    )
    assert resp.status_code == 201, resp.text
    run_dir = Path(resp.json()["run_dir"])
    # No DOCX should be staged for a markdown-only filesystem resume.
    assert not (run_dir / "input" / "master_resume.docx").exists()
    md = (run_dir / "input" / "master_resume.md").read_text(encoding="utf-8")
    assert "SENTINEL_MD" in md

    shutil.rmtree(run_dir, ignore_errors=True)


def test_post_run_with_unknown_filesystem_id_returns_404(client, tmp_path, monkeypatch):
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\n", encoding="utf-8")
    (candidate_root / "project_notes.md").write_text("notes\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text("# prompt\n", encoding="utf-8")
    runs_root = tmp_path / "runs"
    runs_root.mkdir()
    master_resumes_root = tmp_path / "fs"
    master_resumes_root.mkdir()

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))
    monkeypatch.setenv("JOBAPPLY_MASTER_RESUMES_ROOT", str(master_resumes_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build",
        },
    ).json()
    resp = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": "fs:000000000000000a"},
    )
    assert resp.status_code == 404


def test_post_run_with_database_resume_still_writes_master_resume_md(
    client, tmp_path, monkeypatch
):
    """Existing DB-backed flow must continue producing master_resume.md
    without staging master_resume.docx."""
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\n", encoding="utf-8")
    (candidate_root / "project_notes.md").write_text("notes\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text("# prompt\n", encoding="utf-8")
    runs_root = tmp_path / "runs"
    runs_root.mkdir()

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build",
        },
    ).json()
    resume = client.post(
        "/master-resumes",
        json={"name": "main", "content_markdown": "# DB_RESUME_SENTINEL\nbody\n"},
    ).json()
    resp = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": resume["id"]},
    )
    assert resp.status_code == 201, resp.text
    run_dir = Path(resp.json()["run_dir"])

    assert not (run_dir / "input" / "master_resume.docx").exists()
    md = (run_dir / "input" / "master_resume.md").read_text(encoding="utf-8")
    assert "DB_RESUME_SENTINEL" in md

    shutil.rmtree(run_dir, ignore_errors=True)


def test_post_run_creates_row_and_directory(client, tmp_path, monkeypatch):
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\n", encoding="utf-8")
    (candidate_root / "project_notes.md").write_text("notes\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text("# prompt\n", encoding="utf-8")
    runs_root = tmp_path / "runs"
    runs_root.mkdir()

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build things",
        },
    ).json()
    resume = client.post(
        "/master-resumes",
        json={"name": "main", "content_markdown": "# resume\n"},
    ).json()

    resp = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": resume["id"]},
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["status"] == "created"
    assert body["prompt_hash"]
    assert body["input_hash"]
    run_dir = Path(body["run_dir"])
    assert run_dir.is_dir()
    assert (run_dir / "metadata.json").is_file()

    # GET endpoints round-trip
    listed = client.get("/runs").json()
    assert any(r["id"] == body["id"] for r in listed)
    one = client.get(f"/runs/{body['id']}").json()
    assert one["id"] == body["id"]

    # Cleanup the created run dir from outside tmp_path's auto-cleanup scope —
    # in this test it lives under tmp_path, so pytest will handle it.
    shutil.rmtree(run_dir, ignore_errors=True)


def _repo_root() -> Path:
    # backend/tests/test_run_directory.py -> backend/tests -> backend -> repo
    return Path(__file__).resolve().parents[2]


def test_runtime_prompt_requests_docx_skill():
    """The shipped runtime prompt must ask Claude to use the DOCX skill.

    Tracking task 074: the auto path should explicitly request the
    DOCX / Word document skill when generating tailored_resume.docx so
    Claude produces a real Word document instead of plain text dumped
    into a .docx.
    """
    prompt_path = _repo_root() / "runtime_prompts" / "resume_tailoring.md"
    text = prompt_path.read_text(encoding="utf-8")
    assert "DOCX / Word document skill" in text
    # The prompt must say the DOCX is not a plain-text dump.
    assert "not a plain-text dump" in text
    # The prompt must say a source DOCX in input/ should be used as a
    # formatting reference when available.
    assert "source DOCX" in text and "input/" in text
    # The non-interactive contract from task 057 must still be present so
    # the prompt cannot drift back into a conversational response.
    assert "Do not ask clarifying questions" in text
    assert "non-interactive" in text


def test_expected_outputs_includes_ats_audit():
    """Task 092: ats_audit.md is part of the required output contract."""
    from app.run_directory import EXPECTED_OUTPUTS

    assert "ats_audit.md" in EXPECTED_OUTPUTS


def test_runtime_prompt_requests_office_word_mcp():
    """The shipped runtime prompt must prefer the Office Word MCP server.

    Tracking task 075: the auto path should explicitly tell Claude Code
    to use the ``word-document-server`` MCP server when available, and
    to prioritize it over the DOCX skill / fallback generation.
    """
    prompt_path = _repo_root() / "runtime_prompts" / "resume_tailoring.md"
    text = prompt_path.read_text(encoding="utf-8")

    # The prompt must name the connected MCP server explicitly so Claude
    # Code can pick the right tool when it inspects available MCPs.
    assert "word-document-server" in text
    assert "Office Word MCP" in text

    # The Office Word MCP must come before the DOCX skill, and the DOCX
    # skill must come before fallback DOCX generation, in the priority
    # ordering Claude Code is told to follow.
    mcp_pos = text.find("Office Word MCP")
    skill_pos = text.find("DOCX / Word document skill")
    fallback_pos = text.find("fallback DOCX generation")
    assert mcp_pos != -1 and skill_pos != -1 and fallback_pos != -1
    assert mcp_pos < skill_pos < fallback_pos

    # If a source DOCX exists, the prompt must say to copy/edit it as the
    # base rather than rebuilding from scratch.
    assert "copy it as the editable base" in text
    assert "edit relevant text in place" in text

    # The prompt must still say the DOCX is not a plain-text dump.
    assert "not a plain-text dump" in text

    # The prompt must still require post-generation validation that the
    # DOCX exists with nonzero size.
    assert "exists and has nonzero size" in text

    # The non-interactive contract from task 057 must still be present.
    assert "Do not ask clarifying questions" in text
    assert "non-interactive" in text
