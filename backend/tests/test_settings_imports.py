"""File-import endpoints for Settings (task 121).

Master resumes and evidence sources are imported through a file upload
that copies the file into an app-managed ``candidate_context/`` folder,
replacing the old manual ``source_path`` text field. These tests cover
extension validation, filename sanitization, path-traversal defence,
collision handling, and the copy-into-managed-folder contract.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest


def _docx_bytes(body_text: str = "Hello from a real DOCX.") -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe", style="Title")
    doc.add_paragraph("Experience", style="Heading 1")
    doc.add_paragraph(body_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _master_root() -> Path:
    import os

    return Path(os.environ["JOBAPPLY_MASTER_RESUMES_ROOT"])


def _set_candidate_root(monkeypatch, tmp_path: Path) -> Path:
    root = tmp_path / "candidate_context"
    root.mkdir()
    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(root))
    return root


# ---- master resume import ----

def test_master_resume_import_accepts_docx(client):
    data = _docx_bytes("UNIQUE_DOCX_SENTINEL")
    resp = client.post(
        "/master-resumes/import-file",
        files={
            "file": (
                "Calvin Resume.docx",
                data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["source_type"] == "master_resume"
    assert body["source_format"] == "docx"
    assert body["original_filename"] == "Calvin Resume.docx"
    assert body["id"].startswith("fs:")

    # Now discoverable via the master-resume list.
    listed = client.get("/master-resumes").json()
    assert any(e["id"] == body["id"] for e in listed)


def test_master_resume_import_copies_file_into_managed_folder(client):
    resp = client.post(
        "/master-resumes/import-file",
        files={"file": ("resume.md", b"# Resume\nbody\n", "text/markdown")},
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    copied = _master_root() / body["name"]
    assert copied.is_file()
    assert copied.read_text(encoding="utf-8").startswith("# Resume")
    # Stored path is the managed (project-relative) path, not a user path.
    assert body["stored_path"].endswith(body["name"])
    assert ".." not in body["stored_path"]


def test_master_resume_import_rejects_unsupported_extension(client):
    resp = client.post(
        "/master-resumes/import-file",
        files={"file": ("resume.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )
    assert resp.status_code == 400
    assert "unsupported" in resp.json()["detail"].lower()
    # Nothing was written.
    assert list(_master_root().iterdir()) == []


def test_master_resume_import_sanitizes_unsafe_filename(client):
    resp = client.post(
        "/master-resumes/import-file",
        files={"file": ("my résumé (final)!!.md", b"body\n", "text/markdown")},
    )
    assert resp.status_code == 201, resp.text
    name = resp.json()["name"]
    # Only safe filename characters survive; the extension is preserved.
    assert name.endswith(".md")
    assert all(c.isalnum() or c in "._-" for c in name)
    assert (_master_root() / name).is_file()


def test_master_resume_import_prevents_path_traversal(client, tmp_path):
    # A malicious filename must not escape the managed folder.
    resp = client.post(
        "/master-resumes/import-file",
        files={
            "file": ("../../../../tmp/pwned.md", b"evil\n", "text/markdown")
        },
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert ".." not in body["stored_path"]
    # The file lands inside the managed root, named by its basename only.
    assert body["name"] == "pwned.md"
    assert (_master_root() / "pwned.md").is_file()
    # And nothing escaped to the parent directories.
    assert not (_master_root().parent.parent / "pwned.md").exists()


def test_master_resume_import_handles_duplicate_filenames(client):
    first = client.post(
        "/master-resumes/import-file",
        files={"file": ("resume.md", b"one\n", "text/markdown")},
    )
    second = client.post(
        "/master-resumes/import-file",
        files={"file": ("resume.md", b"two\n", "text/markdown")},
    )
    assert first.status_code == 201
    assert second.status_code == 201
    n1 = first.json()["name"]
    n2 = second.json()["name"]
    assert n1 == "resume.md"
    assert n2 == "resume_2.md"
    # Both files exist; the first was not overwritten.
    assert (_master_root() / "resume.md").read_text(encoding="utf-8") == "one\n"
    assert (_master_root() / "resume_2.md").read_text(encoding="utf-8") == "two\n"


# ---- evidence source import ----

def test_evidence_import_accepts_md_and_txt(client, tmp_path, monkeypatch):
    _set_candidate_root(monkeypatch, tmp_path)

    md = client.post(
        "/evidence-sources/import-file",
        files={"file": ("backend.md", b"# Backend\n", "text/markdown")},
    )
    txt = client.post(
        "/evidence-sources/import-file",
        files={"file": ("notes.txt", b"plain notes\n", "text/plain")},
    )
    assert md.status_code == 201, md.text
    assert txt.status_code == 201, txt.text
    assert md.json()["source_format"] == "md"
    assert txt.json()["source_format"] == "txt"
    assert md.json()["source_type"] == "evidence_bank"

    sources = client.get("/evidence-sources").json()
    names = {s["name"] for s in sources}
    assert {"backend.md", "notes.txt"} <= names


def test_evidence_import_copies_file_into_managed_folder(
    client, tmp_path, monkeypatch
):
    candidate_root = _set_candidate_root(monkeypatch, tmp_path)
    resp = client.post(
        "/evidence-sources/import-file",
        files={"file": ("evidence.md", b"EVIDENCE_SENTINEL\n", "text/markdown")},
    )
    assert resp.status_code == 201, resp.text
    copied = candidate_root / "evidence_banks" / resp.json()["name"]
    assert copied.is_file()
    assert "EVIDENCE_SENTINEL" in copied.read_text(encoding="utf-8")
    assert "evidence_banks" in resp.json()["stored_path"]


def test_evidence_import_rejects_unsupported_extension(
    client, tmp_path, monkeypatch
):
    _set_candidate_root(monkeypatch, tmp_path)
    resp = client.post(
        "/evidence-sources/import-file",
        files={"file": ("data.exe", b"MZ", "application/octet-stream")},
    )
    assert resp.status_code == 400
    assert "unsupported" in resp.json()["detail"].lower()


# ---- capture routes still work for the extension (task 121 acceptance) ----

def test_extension_capture_post_still_works(client):
    """The browser extension POSTs to /captures; that must keep working.

    See docs/contracts/browser_extension_capture.md — the extension
    depends on this route, so the capture-inbox cleanup must not remove
    backend capture support.
    """
    resp = client.post(
        "/captures",
        json={
            "source_platform": "linkedin",
            "capture_method": "browser_extension_current_page",
            "external_url": "https://www.linkedin.com/jobs/view/123456789/",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "Build models.",
        },
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["company"] == "Acme"
    # Listing (used by the demoted inbox + activity center) still responds.
    listed = client.get("/captures")
    assert listed.status_code == 200
    assert any(c["id"] == body["id"] for c in listed.json())
