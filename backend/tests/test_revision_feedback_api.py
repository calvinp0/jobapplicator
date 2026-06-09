"""HTTP-layer tests for the revision-feedback endpoint (task 045, ADR-008).

The endpoint at ``POST /resume-versions/{id}/revision-feedback`` performs
three actions atomically from the caller's perspective: insert a
``revision_feedbacks`` row, create a follow-up ``ClaudeRun``, and stage
``input/revision_feedback.md`` into the new run directory. These tests
cover the happy path, 404 on missing source draft, 422 on schema
violation, the join-row FK linking the feedback to the new run, and the
feedback file landing inside the new run's input dir.
"""

from __future__ import annotations

from pathlib import Path


CANDIDATE_FILES = (
    "candidate_profile.md",
    "project_notes.md",
    "skills_inventory.md",
    "tailoring_preferences.md",
    "resume_dos_and_donts.md",
)


# Task 113: a tailoring run now imports a validated resume_suggestions.json,
# so the fake completed-run outputs must include a valid suggestions document
# (and a valid structured resume) for ``/runs/{id}/import`` to succeed.
_MIN_RESUME_JSON = (
    '{"header": {"name": "Test Candidate", "contact_items": []},'
    ' "sections": [{"type": "summary", "heading": "PROFESSIONAL SUMMARY",'
    ' "paragraphs": ["Base summary."]}],'
    ' "metadata": {"target_company": "Acme", "target_job_title": "ML Engineer"}}'
)
_MIN_SUGGESTIONS_JSON = (
    '{"target_company": "Acme", "target_job_title": "ML Engineer",'
    ' "suggestions": []}'
)


def _write_min_outputs(output_dir: Path) -> None:
    """Write the required run outputs, with valid JSON for the JSON artifacts."""
    output_dir.mkdir(parents=True, exist_ok=True)
    for name in (
        "tailored_resume.docx",
        "tailored_resume.md",
        "change_log.md",
        "claim_audit.md",
        "ats_audit.md",
    ):
        (output_dir / name).write_bytes(f"content for {name}\n".encode("utf-8"))
    (output_dir / "tailored_resume.json").write_text(
        _MIN_RESUME_JSON, encoding="utf-8"
    )
    (output_dir / "resume_suggestions.json").write_text(
        _MIN_SUGGESTIONS_JSON, encoding="utf-8"
    )


def _prime_fs(tmp_path: Path, monkeypatch) -> Path:
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text(
        "# Prompt\nRead inputs and write outputs.\n", encoding="utf-8"
    )
    # Revision runs read this prompt via the prompt_harness module when
    # the caller supplies revision_feedback.
    (prompts_root / "resume_revision.md").write_text(
        "# Revision Prompt\nApply user feedback to the current tailored draft.\n",
        encoding="utf-8",
    )
    runs_root = tmp_path / "runs"
    runs_root.mkdir()
    overrides_root = tmp_path / "prompt_overrides"

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))
    monkeypatch.setenv("JOBAPPLY_PROMPT_OVERRIDES_ROOT", str(overrides_root))
    return runs_root


def _seed_draft(client, tmp_path: Path, monkeypatch) -> dict:
    """Create a job + master resume + first-draft ResumeVersion via the API.

    Returns the imported resume_version dict so tests can submit feedback
    against a real prior draft.
    """
    _prime_fs(tmp_path, monkeypatch)

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build things",
        },
    ).json()
    resume = client.post(
        "/master-resumes",
        json={"name": "main", "content_markdown": "# resume\n"},
    ).json()

    run = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": resume["id"]},
    ).json()

    run_dir = Path(run["run_dir"])
    _write_min_outputs(run_dir / "output")

    from app.db import SessionLocal
    from app.models import ClaudeRun

    db = SessionLocal()
    try:
        row = db.get(ClaudeRun, run["id"])
        assert row is not None
        row.status = "completed"
        db.commit()
    finally:
        db.close()

    version = client.post(f"/runs/{run['id']}/import").json()
    return version


def test_revision_feedback_happy_path_creates_row_run_and_file(
    client, tmp_path, monkeypatch
):
    draft = _seed_draft(client, tmp_path, monkeypatch)

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={
            "feedback_markdown": "Soften the executive framing in the summary.",
            "structured_flags": {"too_long": True},
        },
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()

    # The response carries enough info for the frontend to route to the
    # new run: the join-row id, the source draft id, and the follow-up
    # ClaudeRun id.
    assert body["source_resume_version_id"] == draft["id"]
    assert body["followup_claude_run_id"]
    assert body["job_id"] == draft["job_id"]
    assert body["status"] == "created"
    assert body["feedback_markdown"].startswith("Soften")

    # The follow-up run exists and has its run directory populated.
    run = client.get(f"/runs/{body['followup_claude_run_id']}").json()
    assert run["status"] == "created"
    run_dir = Path(run["run_dir"])
    assert run_dir.is_dir()

    feedback_file = run_dir / "input" / "revision_feedback.md"
    assert feedback_file.is_file()
    contents = feedback_file.read_text(encoding="utf-8")
    assert f"source_resume_version_id: {draft['id']}" in contents
    assert "Soften the executive framing" in contents
    assert "too_long" in contents


def test_revision_feedback_creates_exactly_one_record_and_one_run(
    client, tmp_path, monkeypatch
):
    """One submission ⇒ one feedback row, one new ClaudeRun (separate from prior)."""
    draft = _seed_draft(client, tmp_path, monkeypatch)
    prior_run_id = draft["claude_run_id"]

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={"feedback_markdown": "Trim the second bullet."},
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()

    # The new ClaudeRun is distinct from the run that produced the source draft.
    assert body["followup_claude_run_id"] != prior_run_id

    # FK link is populated (ADR-008's "join row" model).
    from app.db import SessionLocal
    from app.models import RevisionFeedback

    db = SessionLocal()
    try:
        rows = db.query(RevisionFeedback).all()
        assert len(rows) == 1
        only = rows[0]
        assert only.source_resume_version_id == draft["id"]
        assert only.followup_claude_run_id == body["followup_claude_run_id"]
        assert only.status == "created"
    finally:
        db.close()

    # Only one extra ClaudeRun was created (the prior + the follow-up).
    runs = client.get("/runs").json()
    assert len(runs) == 2
    follow_up_ids = {r["id"] for r in runs} - {prior_run_id}
    assert follow_up_ids == {body["followup_claude_run_id"]}


def test_revision_feedback_missing_source_draft_returns_404(
    client, tmp_path, monkeypatch
):
    _prime_fs(tmp_path, monkeypatch)

    resp = client.post(
        "/resume-versions/does-not-exist/revision-feedback",
        json={"feedback_markdown": "Anything."},
    )
    assert resp.status_code == 404
    assert "resume version" in resp.json()["detail"].lower()


def test_revision_feedback_empty_body_returns_422(client, tmp_path, monkeypatch):
    draft = _seed_draft(client, tmp_path, monkeypatch)

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={"feedback_markdown": ""},
    )
    assert resp.status_code == 422


def test_revision_feedback_missing_field_returns_422(client, tmp_path, monkeypatch):
    draft = _seed_draft(client, tmp_path, monkeypatch)

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={},
    )
    assert resp.status_code == 422


def test_revision_feedback_followup_run_input_dir_contains_expected_files(
    client, tmp_path, monkeypatch
):
    """The follow-up run dir has the standard input set PLUS revision_feedback.md."""
    draft = _seed_draft(client, tmp_path, monkeypatch)

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={"feedback_markdown": "Please tighten the bullets under Acme."},
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()

    run = client.get(f"/runs/{body['followup_claude_run_id']}").json()
    input_dir = Path(run["run_dir"]) / "input"
    present = {p.name for p in input_dir.iterdir() if p.is_file()}

    expected = {
        "job_description.md",
        "master_resume.md",
        "evidence_bank.md",
        "evidence_sources_index.md",
        "candidate_profile.md",
        "project_notes.md",
        "skills_inventory.md",
        "tailoring_preferences.md",
        "resume_dos_and_donts.md",
        "tailoring_prompt.md",
        # Task 098: every run stages a verbatim snapshot of the
        # effective runtime prompt so the run is reproducible even if
        # the prompt body later changes.
        "prompt_snapshot.md",
        "revision_feedback.md",
        # Task 091: revision runs also stage the prior tailored draft so
        # the worker can revise it rather than regenerate from scratch.
        "current_tailored_resume.md",
        "current_tailored_resume.docx",
    }
    assert present == expected


# --- task 091: revision context provenance ---------------------------------


def _make_docx_file(path: Path, body: str = "Real DOCX body content.") -> None:
    """Write a small valid .docx via python-docx for staging tests."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Header", style="Heading 1")
    doc.add_paragraph(body)
    doc.save(str(path))


def test_revision_feedback_stages_master_resume_markdown(
    client, tmp_path, monkeypatch
):
    """A DB-backed master resume's markdown must reach ``input/master_resume.md``."""
    _prime_fs(tmp_path, monkeypatch)

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build things",
        },
    ).json()
    resume = client.post(
        "/master-resumes",
        json={
            "name": "main",
            "content_markdown": "# Resume\nORIGINAL_RESUME_SENTINEL\n",
        },
    ).json()
    run = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": resume["id"]},
    ).json()
    run_dir = Path(run["run_dir"])
    _write_min_outputs(run_dir / "output")

    from app.db import SessionLocal
    from app.models import ClaudeRun

    db = SessionLocal()
    try:
        row = db.get(ClaudeRun, run["id"])
        row.status = "completed"
        db.commit()
    finally:
        db.close()

    draft = client.post(f"/runs/{run['id']}/import").json()

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={"feedback_markdown": "Tighten the summary."},
    )
    assert resp.status_code == 201, resp.text
    follow_dir = Path(
        client.get(f"/runs/{resp.json()['followup_claude_run_id']}").json()["run_dir"]
    )
    master_md = (follow_dir / "input" / "master_resume.md").read_text(
        encoding="utf-8"
    )
    assert "ORIGINAL_RESUME_SENTINEL" in master_md


def test_revision_feedback_stages_filesystem_master_resume_docx(
    client, tmp_path, monkeypatch
):
    """A filesystem-backed DOCX master resume must stage both .docx and .md.

    This is the exact failure mode described in task 091: prior to the
    fix, ``db.get(MasterResume, fs:...)`` returned None and the endpoint
    raised ``master resume not found for source resume version``.
    """
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text(
        "# Prompt\n", encoding="utf-8"
    )
    (prompts_root / "resume_revision.md").write_text(
        "# Revision Prompt\n", encoding="utf-8"
    )
    runs_root = tmp_path / "runs"
    runs_root.mkdir()
    master_resumes_root = candidate_root / "master_resumes"
    master_resumes_root.mkdir()
    docx_src = master_resumes_root / "calvin.docx"
    _make_docx_file(docx_src, body="FS_DOCX_RESUME_SENTINEL")

    overrides_root = tmp_path / "prompt_overrides"
    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))
    monkeypatch.setenv("JOBAPPLY_MASTER_RESUMES_ROOT", str(master_resumes_root))
    monkeypatch.setenv("JOBAPPLY_PROMPT_OVERRIDES_ROOT", str(overrides_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build",
        },
    ).json()
    fs_entry = next(
        e
        for e in client.get("/master-resumes").json()
        if e["source"] == "filesystem"
    )
    run = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": fs_entry["id"]},
    ).json()
    run_dir = Path(run["run_dir"])
    _write_min_outputs(run_dir / "output")

    from app.db import SessionLocal
    from app.models import ClaudeRun

    db = SessionLocal()
    try:
        row = db.get(ClaudeRun, run["id"])
        row.status = "completed"
        db.commit()
    finally:
        db.close()

    draft = client.post(f"/runs/{run['id']}/import").json()

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={"feedback_markdown": "Reword the intro paragraph."},
    )
    assert resp.status_code == 201, resp.text
    follow_dir = Path(
        client.get(f"/runs/{resp.json()['followup_claude_run_id']}").json()["run_dir"]
    )
    assert (follow_dir / "input" / "master_resume.docx").is_file()
    master_md = (follow_dir / "input" / "master_resume.md").read_text(
        encoding="utf-8"
    )
    assert "FS_DOCX_RESUME_SENTINEL" in master_md


def test_revision_feedback_stages_current_tailored_resume_md_and_docx(
    client, tmp_path, monkeypatch
):
    draft = _seed_draft(client, tmp_path, monkeypatch)

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={"feedback_markdown": "Shorten the bullets under the Acme role."},
    )
    assert resp.status_code == 201, resp.text
    follow_dir = Path(
        client.get(f"/runs/{resp.json()['followup_claude_run_id']}").json()["run_dir"]
    )
    md = (follow_dir / "input" / "current_tailored_resume.md").read_text(
        encoding="utf-8"
    )
    # ``content_markdown`` on the imported ResumeVersion mirrors the
    # tailored_resume.md the worker produced for the source draft.
    assert md == "content for tailored_resume.md\n"
    assert (follow_dir / "input" / "current_tailored_resume.docx").is_file()


def test_revision_feedback_writes_revision_request_file(
    client, tmp_path, monkeypatch
):
    draft = _seed_draft(client, tmp_path, monkeypatch)

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={"feedback_markdown": "REVISION_REQUEST_SENTINEL body text."},
    )
    assert resp.status_code == 201, resp.text
    follow_dir = Path(
        client.get(f"/runs/{resp.json()['followup_claude_run_id']}").json()["run_dir"]
    )
    text = (follow_dir / "input" / "revision_feedback.md").read_text(
        encoding="utf-8"
    )
    assert "REVISION_REQUEST_SENTINEL" in text
    assert f"source_resume_version_id: {draft['id']}" in text


def test_revision_feedback_stages_original_evidence_sources(
    client, tmp_path, monkeypatch
):
    _prime_fs(tmp_path, monkeypatch)

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build",
        },
    ).json()
    resume = client.post(
        "/master-resumes",
        json={"name": "main", "content_markdown": "# resume\n"},
    ).json()
    bank = client.post(
        "/evidence-banks",
        json={
            "name": "primary",
            "content_markdown": "# Bank\nORIGINAL_EVIDENCE_SENTINEL\n",
        },
    ).json()
    run = client.post(
        "/runs",
        json={
            "job_id": job["id"],
            "master_resume_id": resume["id"],
            "evidence_source_ids": [bank["id"]],
        },
    ).json()
    run_dir = Path(run["run_dir"])
    _write_min_outputs(run_dir / "output")

    from app.db import SessionLocal
    from app.models import ClaudeRun

    db = SessionLocal()
    try:
        row = db.get(ClaudeRun, run["id"])
        row.status = "completed"
        db.commit()
    finally:
        db.close()

    draft = client.post(f"/runs/{run['id']}/import").json()

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={"feedback_markdown": "Polish phrasing."},
    )
    assert resp.status_code == 201, resp.text
    follow_dir = Path(
        client.get(f"/runs/{resp.json()['followup_claude_run_id']}").json()["run_dir"]
    )
    evidence_files = list(
        (follow_dir / "input" / "evidence_sources").iterdir()
    )
    assert any(
        "ORIGINAL_EVIDENCE_SENTINEL" in p.read_text(encoding="utf-8")
        for p in evidence_files
        if p.suffix == ".md"
    )
    index = (follow_dir / "input" / "evidence_sources_index.md").read_text(
        encoding="utf-8"
    )
    assert "primary" in index


def test_revision_feedback_accepts_additional_evidence_source_ids(
    client, tmp_path, monkeypatch
):
    draft = _seed_draft(client, tmp_path, monkeypatch)
    extra_bank = client.post(
        "/evidence-banks",
        json={
            "name": "extra",
            "content_markdown": "# Extra\nEXTRA_EVIDENCE_SENTINEL\n",
        },
    ).json()

    resp = client.post(
        f"/resume-versions/{draft['id']}/revision-feedback",
        json={
            "feedback_markdown": "Pull in the new project context.",
            "additional_evidence_source_ids": [extra_bank["id"]],
        },
    )
    assert resp.status_code == 201, resp.text
    follow_dir = Path(
        client.get(f"/runs/{resp.json()['followup_claude_run_id']}").json()["run_dir"]
    )
    evidence_files = list(
        (follow_dir / "input" / "evidence_sources").iterdir()
    )
    assert any(
        "EXTRA_EVIDENCE_SENTINEL" in p.read_text(encoding="utf-8")
        for p in evidence_files
        if p.suffix == ".md"
    )
    # The revision_feedback.md frontmatter records the additional ids so
    # the worker can distinguish them from the original sources.
    feedback_text = (follow_dir / "input" / "revision_feedback.md").read_text(
        encoding="utf-8"
    )
    assert "additional_evidence_source_ids" in feedback_text
    assert extra_bank["id"] in feedback_text


def test_revision_feedback_creates_new_resume_version_not_overwriting_source(
    client, tmp_path, monkeypatch
):
    """Importing the follow-up run produces a new version, leaving the source intact."""
    draft = _seed_draft(client, tmp_path, monkeypatch)
    source_id = draft["id"]

    resp = client.post(
        f"/resume-versions/{source_id}/revision-feedback",
        json={"feedback_markdown": "Refine the bullets."},
    )
    assert resp.status_code == 201, resp.text
    followup_run_id = resp.json()["followup_claude_run_id"]
    follow_dir = Path(client.get(f"/runs/{followup_run_id}").json()["run_dir"])
    _write_min_outputs(follow_dir / "output")

    from app.db import SessionLocal
    from app.models import ClaudeRun

    db = SessionLocal()
    try:
        row = db.get(ClaudeRun, followup_run_id)
        row.status = "completed"
        db.commit()
    finally:
        db.close()

    new_version = client.post(f"/runs/{followup_run_id}/import").json()
    assert new_version["id"] != source_id
    assert new_version["version_number"] == draft["version_number"] + 1

    # The source draft row is still present and unchanged.
    fetched_source = client.get(f"/resume-versions/{source_id}").json()
    assert fetched_source["id"] == source_id
    assert fetched_source["content_markdown"] == draft["content_markdown"]


def test_revision_feedback_structured_error_when_master_resume_missing(
    client, tmp_path, monkeypatch
):
    """An unresolvable master resume id returns a structured 422 detail, not a 500."""
    _prime_fs(tmp_path, monkeypatch)

    # Hand-craft a ResumeVersion whose master_resume_id points at nothing
    # — neither a DB row nor a filesystem discovery.
    from app.db import SessionLocal
    from app.models import ClaudeRun, Job, ResumeVersion

    db = SessionLocal()
    try:
        job = Job(
            source_platform="linkedin",
            company="Acme",
            title="ML",
            description_text="...",
        )
        db.add(job)
        db.flush()
        # Synthetic fs id that resolve_filesystem_master_resume will not find.
        run = ClaudeRun(
            job_id=job.id,
            master_resume_id="fs:0000000000000000",
            run_dir=str(tmp_path / "runs" / "orphan"),
            status="completed",
        )
        db.add(run)
        db.flush()
        version = ResumeVersion(
            job_id=job.id,
            master_resume_id="fs:0000000000000000",
            claude_run_id=run.id,
            version_number=1,
            content_markdown="# stub\n",
            source="claude_run",
        )
        db.add(version)
        db.commit()
        version_id = version.id
    finally:
        db.close()

    resp = client.post(
        f"/resume-versions/{version_id}/revision-feedback",
        json={"feedback_markdown": "Anything."},
    )
    assert resp.status_code == 422, resp.text
    detail = resp.json()["detail"]
    assert isinstance(detail, dict)
    assert detail["error"] == "revision_missing_master_resume"
    assert detail["source_resume_version_id"] == version_id
    assert "master resume" in detail["message"].lower()


def test_revision_feedback_legacy_draft_without_claude_run_falls_back(
    client, tmp_path, monkeypatch
):
    """A draft with no ``claude_run_id`` still produces a working revision run."""
    _prime_fs(tmp_path, monkeypatch)

    from app.db import SessionLocal
    from app.models import Job, MasterResume, ResumeVersion

    db = SessionLocal()
    try:
        job = Job(
            source_platform="linkedin",
            company="Acme",
            title="ML",
            description_text="...",
        )
        master = MasterResume(name="legacy", content_markdown="# legacy resume\n")
        db.add_all([job, master])
        db.flush()
        version = ResumeVersion(
            job_id=job.id,
            master_resume_id=master.id,
            claude_run_id=None,
            version_number=1,
            content_markdown="# legacy draft\n",
            source="manual",
        )
        db.add(version)
        db.commit()
        version_id = version.id
    finally:
        db.close()

    resp = client.post(
        f"/resume-versions/{version_id}/revision-feedback",
        json={"feedback_markdown": "Tighten."},
    )
    assert resp.status_code == 201, resp.text
    follow_dir = Path(
        client.get(f"/runs/{resp.json()['followup_claude_run_id']}").json()["run_dir"]
    )
    # No prior evidence sources available, but the index file is still present.
    assert (follow_dir / "input" / "evidence_sources_index.md").is_file()
    assert (follow_dir / "input" / "current_tailored_resume.md").is_file()


def test_runtime_prompt_includes_revision_context():
    """The shipped prompt names the current tailored draft and revision request."""
    prompt = (
        Path(__file__).resolve().parents[2]
        / "runtime_prompts"
        / "resume_tailoring.md"
    ).read_text(encoding="utf-8")
    # Master resume, current tailored draft, evidence sources, and revision
    # request must all be referenced in the prompt body.
    assert "input/master_resume.md" in prompt
    assert "input/current_tailored_resume.md" in prompt
    assert "input/evidence_sources_index.md" in prompt
    assert "input/revision_feedback.md" in prompt


def test_revision_prompt_requests_updated_recruiter_review():
    """Task 108: revision runs must refresh output/recruiter_review.md
    after applying the user's feedback."""
    prompt = (
        Path(__file__).resolve().parents[2]
        / "runtime_prompts"
        / "resume_revision.md"
    ).read_text(encoding="utf-8")
    assert "output/recruiter_review.md" in prompt
    # The instruction wording must explicitly tell Claude to re-review
    # the result after applying the revision request.
    normalized = " ".join(prompt.split()).lower()
    assert (
        "re-review the result as a recruiter/hiring manager and update"
        in normalized
    )
