from __future__ import annotations

from pathlib import Path


def _make_docx(path: Path, *, body_text: str = "Hello world from a real DOCX.") -> None:
    """Write a real DOCX file at ``path`` using python-docx.

    Mirrors the shape used by ``test_docx_extract.py`` so we exercise the
    same code path the worker would when extracting a discovered resume.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe", style="Title")
    doc.add_paragraph("Experience", style="Heading 1")
    doc.add_paragraph(body_text)
    doc.save(str(path))


def _set_master_resumes_root(monkeypatch, tmp_path: Path) -> Path:
    """Point the discovery layer at an isolated temp directory."""
    root = tmp_path / "master_resumes"
    root.mkdir()
    monkeypatch.setenv("JOBAPPLY_MASTER_RESUMES_ROOT", str(root))
    return root


def test_create_master_resume(client):
    r = client.post(
        "/master-resumes",
        json={"name": "main", "content_markdown": "# Resume\n"},
    )
    assert r.status_code == 201, r.text
    body = r.json()
    assert body["name"] == "main"
    assert body["content_markdown"].startswith("# Resume")

    listed = client.get("/master-resumes").json()
    assert len(listed) == 1
    assert listed[0]["id"] == body["id"]
    assert listed[0]["source"] == "database"
    assert listed[0]["is_demo"] is False


def test_filesystem_docx_resume_is_discovered(client, tmp_path, monkeypatch):
    root = _set_master_resumes_root(monkeypatch, tmp_path)
    _make_docx(root / "calvin_resume.docx")

    listed = client.get("/master-resumes").json()
    assert len(listed) == 1
    entry = listed[0]
    assert entry["id"].startswith("fs:")
    assert entry["name"] == "calvin_resume.docx"
    assert entry["source"] == "filesystem"
    assert entry["source_format"] == "docx"
    assert entry["source_path"].endswith("calvin_resume.docx")
    assert entry["is_demo"] is False


def test_filesystem_md_resume_is_discovered(client, tmp_path, monkeypatch):
    root = _set_master_resumes_root(monkeypatch, tmp_path)
    (root / "industry_ml.md").write_text("# ML Resume\nbody\n", encoding="utf-8")

    listed = client.get("/master-resumes").json()
    assert [e["name"] for e in listed] == ["industry_ml.md"]
    assert listed[0]["source_format"] == "md"


def test_filesystem_txt_resume_is_discovered(client, tmp_path, monkeypatch):
    root = _set_master_resumes_root(monkeypatch, tmp_path)
    (root / "plain_resume.txt").write_text("plain resume text\n", encoding="utf-8")

    listed = client.get("/master-resumes").json()
    assert [e["name"] for e in listed] == ["plain_resume.txt"]
    assert listed[0]["source_format"] == "txt"


def test_unsupported_extensions_are_ignored(client, tmp_path, monkeypatch):
    root = _set_master_resumes_root(monkeypatch, tmp_path)
    (root / "resume.pdf").write_bytes(b"%PDF-1.4 fake")
    (root / "resume.rtf").write_text("rtf body", encoding="utf-8")
    (root / "valid.md").write_text("ok\n", encoding="utf-8")

    listed = client.get("/master-resumes").json()
    names = [e["name"] for e in listed]
    assert names == ["valid.md"]


def test_word_lock_and_hidden_files_are_ignored(client, tmp_path, monkeypatch):
    root = _set_master_resumes_root(monkeypatch, tmp_path)
    (root / ".~lock.resume.docx#").write_text("locked", encoding="utf-8")
    (root / "~$resume.docx").write_text("temp", encoding="utf-8")
    (root / ".DS_Store").write_text("mac junk", encoding="utf-8")
    (root / "._appledouble.md").write_text("meta", encoding="utf-8")
    (root / "real.md").write_text("real body\n", encoding="utf-8")

    listed = client.get("/master-resumes").json()
    assert [e["name"] for e in listed] == ["real.md"]


def test_filesystem_ids_are_stable_across_calls(client, tmp_path, monkeypatch):
    root = _set_master_resumes_root(monkeypatch, tmp_path)
    (root / "a.md").write_text("a", encoding="utf-8")
    (root / "b.md").write_text("b", encoding="utf-8")

    first = client.get("/master-resumes").json()
    second = client.get("/master-resumes").json()

    by_name_first = {e["name"]: e["id"] for e in first}
    by_name_second = {e["name"]: e["id"] for e in second}
    assert by_name_first == by_name_second
    assert by_name_first["a.md"] != by_name_first["b.md"]


def test_filesystem_resumes_sort_before_db_resumes(client, tmp_path, monkeypatch):
    """Real files should outrank seeded demo data in the selector."""
    root = _set_master_resumes_root(monkeypatch, tmp_path)
    (root / "real.md").write_text("real\n", encoding="utf-8")

    # Seed a DB-backed resume (mimics the demo seed).
    client.post(
        "/master-resumes",
        json={"name": "Demo Master Resume", "content_markdown": "demo\n"},
    ).raise_for_status()

    listed = client.get("/master-resumes").json()
    assert len(listed) == 2
    # Filesystem entry comes first.
    assert listed[0]["source"] == "filesystem"
    assert listed[0]["name"] == "real.md"
    # Demo entry last, flagged as demo so the UI can de-emphasize it.
    assert listed[1]["source"] == "database"
    assert listed[1]["is_demo"] is True


def test_db_demo_resume_remains_available_when_no_files(client, tmp_path, monkeypatch):
    _set_master_resumes_root(monkeypatch, tmp_path)
    client.post(
        "/master-resumes",
        json={"name": "Demo Master Resume", "content_markdown": "demo\n"},
    ).raise_for_status()

    listed = client.get("/master-resumes").json()
    assert len(listed) == 1
    assert listed[0]["is_demo"] is True


def test_get_filesystem_master_resume_returns_extracted_content(
    client, tmp_path, monkeypatch
):
    root = _set_master_resumes_root(monkeypatch, tmp_path)
    _make_docx(root / "calvin.docx", body_text="UNIQUE_DOCX_SENTINEL_12345")

    listed = client.get("/master-resumes").json()
    fs_id = listed[0]["id"]

    detail = client.get(f"/master-resumes/{fs_id}").json()
    assert detail["id"] == fs_id
    assert detail["source"] == "filesystem"
    assert detail["source_format"] == "docx"
    assert "UNIQUE_DOCX_SENTINEL_12345" in detail["content_markdown"]


def test_get_unknown_filesystem_id_returns_404(client, tmp_path, monkeypatch):
    _set_master_resumes_root(monkeypatch, tmp_path)
    r = client.get("/master-resumes/fs:deadbeefcafef00d")
    assert r.status_code == 404
