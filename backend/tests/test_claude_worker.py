from __future__ import annotations

import os
import shutil
import sys
import textwrap
from pathlib import Path

import pytest


CANDIDATE_FILES = (
    "candidate_profile.md",
    "project_notes.md",
    "skills_inventory.md",
    "tailoring_preferences.md",
    "resume_dos_and_donts.md",
)


ALL_OUTPUTS = (
    "tailored_resume.md",
    "tailored_resume.docx",
    "tailored_resume.json",
    "resume_suggestions.json",
    "change_log.md",
    "claim_audit.md",
    "ats_audit.md",
)


MINIMAL_VALID_SUGGESTIONS_JSON = (
    '{\n'
    '  "target_company": "Acme",\n'
    '  "target_job_title": "ML Engineer",\n'
    '  "suggestions": [\n'
    '    {"id": "sug_001", "section_id": "professional_summary",\n'
    '     "section_heading": "PROFESSIONAL SUMMARY",\n'
    '     "operation": "replace_section_text",\n'
    '     "current_text": "old", "suggested_text": "new",\n'
    '     "reason": "Aligns with the role.", "risk": "low",\n'
    '     "confidence": 0.8, "status": "pending"}\n'
    '  ]\n'
    '}\n'
)


MINIMAL_VALID_RESUME_JSON = (
    '{\n'
    '  "header": {"name": "Test Candidate", "contact_items": ["test@example.com"]},\n'
    '  "sections": [\n'
    '    {"type": "summary", "heading": "SUMMARY",\n'
    '     "paragraphs": ["Engineer with experience building things."]},\n'
    '    {"type": "experience", "heading": "EXPERIENCE",\n'
    '     "entries": [{"title": "Engineer", "organization": "Acme",\n'
    '                  "dates": "2024", "bullets": ["Built systems."]}]}\n'
    '  ],\n'
    '  "metadata": {"target_company": "Acme", "target_job_title": "ML Engineer"}\n'
    '}\n'
)


def _write_fake_binary(
    tmp_path: Path,
    *,
    exit_code: int = 0,
    write_outputs: tuple[str, ...] = ALL_OUTPUTS,
    extra_body: str = "",
) -> Path:
    """Write an executable Python script that imitates Claude Code.

    Writes a marker file (``where_am_i.txt``) recording the subprocess cwd,
    drains and records stdin to ``stdin_received.txt`` so tests can assert the
    worker piped the prompt contents in non-interactive mode, writes the
    requested subset of ``output/`` files, then exits with ``exit_code``.
    """
    binary = tmp_path / f"fake_claude_{exit_code}_{'_'.join(write_outputs) or 'none'}"
    outputs_repr = repr(list(write_outputs))
    json_payload_repr = repr(MINIMAL_VALID_RESUME_JSON)
    suggestions_payload_repr = repr(MINIMAL_VALID_SUGGESTIONS_JSON)
    body = textwrap.dedent(
        f"""\
        #!{sys.executable}
        import os
        import sys
        from pathlib import Path

        cwd = Path.cwd()
        print(f"fake claude running in {{cwd}}", flush=True)
        print(f"argv={{sys.argv!r}}", flush=True)
        (cwd / "where_am_i.txt").write_text(str(cwd) + "\\n", encoding="utf-8")
        try:
            stdin_blob = sys.stdin.read()
        except Exception:
            stdin_blob = ""
        (cwd / "stdin_received.txt").write_text(stdin_blob, encoding="utf-8")
        out = cwd / "output"
        out.mkdir(parents=True, exist_ok=True)
        _json_payload = {json_payload_repr}
        _suggestions_payload = {suggestions_payload_repr}
        for name in {outputs_repr}:
            if name == "tailored_resume.json":
                (out / name).write_text(_json_payload, encoding="utf-8")
            elif name == "resume_suggestions.json":
                (out / name).write_text(_suggestions_payload, encoding="utf-8")
            else:
                (out / name).write_bytes(f"content for {{name}}\\n".encode("utf-8"))
        {extra_body}
        sys.exit({exit_code})
        """
    )
    binary.write_text(body, encoding="utf-8")
    binary.chmod(0o755)
    return binary


def _seed_run(client, tmp_path: Path, monkeypatch) -> dict:
    """Create the supporting files and POST /runs to get a real run row."""
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\nbody\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text(
        "# Prompt\n"
        "You are running inside a non-interactive backend job.\n"
        "Do not ask clarifying questions.\n"
        "Do the thing.\n",
        encoding="utf-8",
    )
    runs_root = tmp_path / "runs"
    runs_root.mkdir()

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build things",
        },
    ).json()
    resume = client.post(
        "/master-resumes",
        json={"name": "main", "content_markdown": "# resume\n"},
    ).json()
    resp = client.post(
        "/runs",
        json={"job_id": job["id"], "master_resume_id": resume["id"]},
    )
    assert resp.status_code == 201, resp.text
    return resp.json()


def test_invoke_run_success_transitions_status(client, tmp_path, monkeypatch):
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    assert run["status"] == "created"
    assert run["started_at"] is None
    assert run["completed_at"] is None

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    assert body["status"] == "completed"
    assert body["started_at"] is not None
    assert body["completed_at"] is not None
    assert body["error_message"] is None

    run_dir = Path(body["run_dir"])
    log_path = run_dir / "run.log"
    assert log_path.is_file()
    assert log_path.stat().st_size > 0
    log_text = log_path.read_text(encoding="utf-8")
    assert "fake claude running in" in log_text

    # All required output files Claude was supposed to write.
    for name in ALL_OUTPUTS:
        assert (run_dir / "output" / name).is_file()


def test_invoke_run_failure_records_error(client, tmp_path, monkeypatch):
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=3)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    assert body["status"] == "failed"
    assert body["started_at"] is not None
    assert body["completed_at"] is not None
    assert body["error_message"]
    assert "3" in body["error_message"]

    run_dir = Path(body["run_dir"])
    assert (run_dir / "run.log").stat().st_size > 0


def test_invoke_run_subprocess_cwd_is_run_dir(client, tmp_path, monkeypatch):
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    run_dir = Path(body["run_dir"]).resolve()
    marker = run_dir / "where_am_i.txt"
    assert marker.is_file()
    recorded = Path(marker.read_text(encoding="utf-8").strip()).resolve()
    assert recorded == run_dir
    # Make sure the subprocess did NOT run from the project root or backend dir.
    assert recorded != Path.cwd().resolve()


def test_invoke_run_missing_binary_marks_failed(client, tmp_path, monkeypatch):
    run = _seed_run(client, tmp_path, monkeypatch)
    monkeypatch.setenv(
        "JOBAPPLY_CLAUDE_BINARY", str(tmp_path / "definitely_not_a_binary")
    )

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    assert body["status"] == "failed"
    assert body["error_message"]
    # After ADR-009 the worker quotes the provider id, not the literal
    # string "claude", so a future provider-specific failure stays
    # distinguishable in logs.
    assert "binary not found" in body["error_message"]
    assert "'claude_code'" in body["error_message"]


def test_invoke_run_unknown_id_returns_404(client, tmp_path, monkeypatch):
    # Need the env to be valid enough for the route to import, though it
    # short-circuits before touching the subprocess.
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post("/runs/does-not-exist/invoke")
    assert resp.status_code == 404
    assert "claude run" in resp.json()["detail"].lower()


def test_invoke_run_dry_run_skips_subprocess(client, tmp_path, monkeypatch):
    run = _seed_run(client, tmp_path, monkeypatch)
    # Binary points at something that would fail if actually executed.
    monkeypatch.setenv(
        "JOBAPPLY_CLAUDE_BINARY", str(tmp_path / "would_explode_if_used")
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_DRY_RUN", "1")

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    assert body["status"] == "completed"
    assert body["started_at"] is not None
    assert body["completed_at"] is not None
    assert body["error_message"] is None

    run_dir = Path(body["run_dir"])
    log_text = (run_dir / "run.log").read_text(encoding="utf-8")
    assert "dry-run" in log_text

    # Dry-run must populate the full output contract so the run is importable.
    for name in ALL_OUTPUTS:
        assert (run_dir / "output" / name).is_file(), f"missing dry-run output: {name}"

    # The placeholder docx must be a real zip (Word package) so the open-file
    # flow can hand it to the host application without a corrupt-file error.
    import zipfile

    docx = run_dir / "output" / "tailored_resume.docx"
    assert zipfile.is_zipfile(docx)
    with zipfile.ZipFile(docx) as zf:
        names = set(zf.namelist())
    assert {"[Content_Types].xml", "_rels/.rels", "word/document.xml"} <= names


def test_invoke_run_zero_exit_with_missing_outputs_marks_failed(
    client, tmp_path, monkeypatch
):
    """Claude exit code 0 is not enough — the output contract must be satisfied.

    Task 111 made ``output/tailored_resume.json`` the first gated output:
    the deterministic renderer needs it before it can produce the final
    DOCX, so a Claude that exits cleanly without writing the structured
    JSON fails the run with a clear missing-file message.
    """
    run = _seed_run(client, tmp_path, monkeypatch)
    # Subprocess "succeeds" but writes the markdown projection without
    # producing the structured JSON the renderer needs.
    binary = _write_fake_binary(
        tmp_path,
        exit_code=0,
        write_outputs=("tailored_resume.md",),
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    assert body["status"] == "failed"
    assert body["completed_at"] is not None
    message = body["error_message"]
    assert "expected output file missing" in message
    assert "output/tailored_resume.json" in message


def test_invoke_run_zero_exit_with_json_only_marks_failed_for_remaining_outputs(
    client, tmp_path, monkeypatch
):
    """JSON present + everything else missing still fails the contract.

    The deterministic renderer produces the DOCX and the template
    fidelity audit from the JSON, so those two files appear; the rest
    of the output contract (markdown, change log, audits) is still
    required.
    """
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(
        tmp_path,
        exit_code=0,
        write_outputs=("tailored_resume.json",),
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    assert body["status"] == "failed"
    message = body["error_message"]
    assert "expected output file missing" in message
    for missing in (
        "output/tailored_resume.md",
        "output/change_log.md",
        "output/claim_audit.md",
        "output/ats_audit.md",
    ):
        assert missing in message, f"{missing!r} not in {message!r}"
    # JSON and DOCX (rendered from JSON) must not appear as missing.
    assert "output/tailored_resume.json" not in message
    assert "output/tailored_resume.docx" not in message


def test_invoke_run_zero_exit_with_all_outputs_marks_completed(
    client, tmp_path, monkeypatch
):
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=ALL_OUTPUTS)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"
    assert body["error_message"] is None


def test_invoke_run_missing_prompt_marks_failed(client, tmp_path, monkeypatch):
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    prompt_path = Path(run["run_dir"]) / "input" / "tailoring_prompt.md"
    prompt_path.unlink()

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "failed"
    assert "tailoring prompt not found" in body["error_message"]


def test_invoke_run_passes_default_permission_mode(client, tmp_path, monkeypatch):
    """Worker must pass ``--permission-mode acceptEdits`` so writes auto-approve."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))
    monkeypatch.delenv("JOBAPPLY_CLAUDE_PERMISSION_MODE", raising=False)

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    # The fake binary echoes its argv into the log via stdout.
    assert "--permission-mode" in log_text
    assert "acceptEdits" in log_text
    # The worker must use non-interactive (--print) mode so Claude doesn't
    # open a REPL.
    assert "--print" in log_text
    # Worker-owned progress lines record the non-interactive launch context.
    assert "jobapply: launching Claude Code in non-interactive mode" in log_text
    assert "jobapply: prompt file=input/tailoring_prompt.md" in log_text
    assert "jobapply: launching Claude Code with cwd=" in log_text
    assert "jobapply: permission mode=acceptEdits" in log_text
    assert "jobapply: output directory=" in log_text


def test_invoke_run_logs_word_docx_tooling_requested(client, tmp_path, monkeypatch):
    """Worker must log that Word/DOCX tooling was requested for DOCX output.

    Tracking task 075: the runtime prompt asks Claude Code to prefer the
    Office Word MCP server (``word-document-server``), then the DOCX /
    Word document skill, then fall back to existing generation. The
    worker has no reliable cross-version way to detect MCP or skill
    installation, so it logs that each was requested and leaves
    availability unknown.
    """
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "jobapply: Word/DOCX tooling requested for DOCX generation" in log_text
    assert "jobapply: Office Word MCP server requested if available" in log_text
    assert "jobapply: DOCX skill requested if available" in log_text
    # The worker should be honest that detection is not implemented rather
    # than claiming either capability is definitely installed.
    assert "jobapply: Office Word MCP availability unknown" in log_text
    assert "jobapply: DOCX skill availability unknown" in log_text


def test_invoke_run_pipes_prompt_contents_via_stdin(client, tmp_path, monkeypatch):
    """Worker must pipe the prompt file contents to Claude, not the file path.

    The previous (broken) invocation passed the prompt file path as a
    conversational argument, which made Claude respond with "Let me know
    which direction" instead of executing the contract. Verify the worker
    now hands Claude the prompt body via stdin, and that the body still
    carries the non-interactive instructions from the runtime prompt.
    """
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    run_dir = Path(body["run_dir"])
    stdin_blob = (run_dir / "stdin_received.txt").read_text(encoding="utf-8")

    # The prompt body, not just the path, must reach the subprocess.
    prompt_text = (run_dir / "input" / "tailoring_prompt.md").read_text(
        encoding="utf-8"
    )
    assert stdin_blob == prompt_text
    # And the prompt body must carry the non-interactive contract so Claude
    # cannot drift back into "what do you want me to do?" mode.
    assert "non-interactive" in stdin_blob
    assert "Do not ask clarifying questions" in stdin_blob

    # The argv must NOT include the bare prompt path positional that caused
    # the original conversational drift.
    log_text = (run_dir / "run.log").read_text(encoding="utf-8")
    assert "'input/tailoring_prompt.md'" not in log_text


def test_invoke_run_missing_outputs_marks_failed_under_print_mode(
    client, tmp_path, monkeypatch
):
    """Non-interactive launch must still enforce the output contract.

    A fake Claude that writes nothing (mimicking a run that drifted into a
    conversational response and produced no files) must result in
    ``failed`` even though the subprocess exits cleanly. Task 111 made
    the structured JSON the first gate, so the missing-file message
    names it explicitly.
    """
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=())
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "failed"
    message = body["error_message"]
    assert "expected output file missing" in message
    assert "output/tailored_resume.json" in message


def test_invoke_run_permission_mode_env_override(client, tmp_path, monkeypatch):
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))
    monkeypatch.setenv("JOBAPPLY_CLAUDE_PERMISSION_MODE", "plan")

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "--permission-mode plan" in log_text
    assert "jobapply: permission mode=plan" in log_text
    assert "jobapply: permission mode=acceptEdits" not in log_text


def test_invoke_run_extra_args_still_passed(client, tmp_path, monkeypatch):
    """Existing JOBAPPLY_CLAUDE_EXTRA_ARGS contract is preserved alongside permission-mode."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))
    monkeypatch.setenv("JOBAPPLY_CLAUDE_EXTRA_ARGS", "--verbose")

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "--permission-mode" in log_text
    assert "--verbose" in log_text


def test_invoke_run_creates_output_dir_if_missing(client, tmp_path, monkeypatch):
    """Worker must (re-)create output/ before launching so writes don't fail."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    output_dir = Path(run["run_dir"]) / "output"
    assert output_dir.is_dir()
    shutil.rmtree(output_dir)
    assert not output_dir.exists()

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed", body
    assert output_dir.is_dir()
    for name in ALL_OUTPUTS:
        assert (output_dir / name).is_file()


def _seed_run_with_provider(
    client, tmp_path, monkeypatch, provider_id: str
) -> dict:
    """Variant of ``_seed_run`` that POSTs an explicit ``llm_provider``."""
    candidate_root = tmp_path / "candidate_context"
    candidate_root.mkdir()
    for name in CANDIDATE_FILES:
        (candidate_root / name).write_text(f"# {name}\nbody\n", encoding="utf-8")
    prompts_root = tmp_path / "runtime_prompts"
    prompts_root.mkdir()
    (prompts_root / "resume_tailoring.md").write_text("# prompt\n", encoding="utf-8")
    runs_root = tmp_path / "runs"
    runs_root.mkdir()

    monkeypatch.setenv("JOBAPPLY_CANDIDATE_CONTEXT_ROOT", str(candidate_root))
    monkeypatch.setenv("JOBAPPLY_RUNTIME_PROMPTS_ROOT", str(prompts_root))
    monkeypatch.setenv("JOBAPPLY_RUNS_ROOT", str(runs_root))

    job = client.post(
        "/jobs",
        json={
            "source_platform": "linkedin",
            "company": "Acme",
            "title": "ML Engineer",
            "description_text": "build things",
        },
    ).json()
    resume = client.post(
        "/master-resumes",
        json={"name": "main", "content_markdown": "# resume\n"},
    ).json()
    resp = client.post(
        "/runs",
        json={
            "job_id": job["id"],
            "master_resume_id": resume["id"],
            "llm_provider": provider_id,
        },
    )
    assert resp.status_code == 201, resp.text
    return resp.json()


def test_invoke_run_dry_run_preserves_codex_provider(
    client, tmp_path, monkeypatch
):
    """Per ADR-009 the persisted provider id survives a dry-run round trip.

    Dry-run skips the subprocess so we can verify the registry plumbing
    end-to-end without depending on a fake binary: the run is created
    with ``llm_provider=codex``, the column is persisted, ``metadata.json``
    records the same value, and dry-run still marks the run completed
    (providers are not consulted in dry-run mode).
    """
    import json

    run = _seed_run_with_provider(client, tmp_path, monkeypatch, "codex")
    assert run["llm_provider"] == "codex"

    metadata = json.loads(
        (Path(run["run_dir"]) / "metadata.json").read_text(encoding="utf-8")
    )
    assert metadata["llm_provider"] == "codex"

    monkeypatch.setenv("JOBAPPLY_CLAUDE_DRY_RUN", "1")
    monkeypatch.setenv(
        "JOBAPPLY_CODEX_BINARY", str(tmp_path / "would_explode_if_used")
    )
    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()

    assert body["status"] == "completed"
    assert body["llm_provider"] == "codex"
    # Re-fetch to confirm the column is read back as codex after the
    # subprocess-less dry-run completes (the value is set at create-time
    # and must not be mutated by invoke).
    fetched = client.get(f"/runs/{run['id']}").json()
    assert fetched["llm_provider"] == "codex"


def test_invoke_does_not_mutate_resume_versions(client, tmp_path, monkeypatch):
    """The worker only records run lifecycle; it must not create ResumeVersion rows."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    # output_hash stays empty — that's task 008's job.
    assert body["output_hash"] is None


def _write_master_resume_docx(input_dir: Path) -> Path:
    """Create a small but valid master_resume.docx under ``input_dir``.

    Generated inline via python-docx so tests do not depend on a
    committed binary fixture.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe", style="Title")
    doc.add_paragraph("Experience", style="Heading 1")
    doc.add_paragraph("Acme Corp — Staff Engineer", style="Heading 2")
    doc.add_paragraph("Built distributed systems.", style="List Bullet")
    path = input_dir / "master_resume.docx"
    doc.save(str(path))
    return path


def test_invoke_run_detects_and_extracts_master_resume_docx(
    client, tmp_path, monkeypatch
):
    """When a source DOCX is staged in input/, the worker must extract it
    before launching Claude and emit the documented run-log lines."""
    run = _seed_run(client, tmp_path, monkeypatch)
    input_dir = Path(run["run_dir"]) / "input"
    _write_master_resume_docx(input_dir)

    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    extracted = input_dir / "master_resume_extracted.md"
    assert extracted.is_file()
    extracted_text = extracted.read_text(encoding="utf-8")
    # Filename of the source DOCX must be recorded in the extracted markdown.
    assert "Source DOCX: input/master_resume.docx" in extracted_text
    assert "Jane Doe" in extracted_text

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "jobapply: checking for master resume DOCX" in log_text
    assert (
        "jobapply: found source resume DOCX=input/master_resume.docx"
        in log_text
    )
    assert (
        "jobapply: extracted source resume DOCX to "
        "input/master_resume_extracted.md"
    ) in log_text


def test_invoke_run_markdown_only_resume_still_runs(
    client, tmp_path, monkeypatch
):
    """Existing markdown-only flow must keep working — no DOCX, no
    extracted markdown, but the worker still completes the four-output
    contract and logs that it checked for a DOCX."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    input_dir = Path(run["run_dir"]) / "input"
    assert not (input_dir / "master_resume.docx").exists()

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    # No DOCX, so no extracted file should be created.
    assert not (input_dir / "master_resume_extracted.md").exists()
    assert not (input_dir / "master_resume_extraction_error.md").exists()

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "jobapply: checking for master resume DOCX" in log_text
    # The "found" line must NOT appear when no DOCX is present.
    assert "jobapply: found source resume DOCX" not in log_text

    # Output contract still satisfied.
    for name in ALL_OUTPUTS:
        assert (Path(body["run_dir"]) / "output" / name).is_file()


def test_invoke_run_extraction_failure_logs_and_continues_with_markdown(
    client, tmp_path, monkeypatch
):
    """A DOCX that python-docx cannot open must produce an error file and
    a failure log line — but with a markdown resume present, the run
    proceeds and Claude still runs against the markdown evidence."""
    run = _seed_run(client, tmp_path, monkeypatch)
    input_dir = Path(run["run_dir"]) / "input"
    # Stage an invalid DOCX (master_resume.md already exists from _seed_run).
    (input_dir / "master_resume.docx").write_bytes(b"not a real docx")

    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    # markdown fallback exists, so the run still completes.
    assert body["status"] == "completed", body

    error_file = input_dir / "master_resume_extraction_error.md"
    assert error_file.is_file()
    assert "Source DOCX: input/master_resume.docx" in error_file.read_text(
        encoding="utf-8"
    )

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "jobapply: failed to extract source resume DOCX" in log_text


def test_invoke_run_extraction_failure_without_markdown_marks_failed(
    client, tmp_path, monkeypatch
):
    """If extraction fails and no markdown resume is on disk, the run
    must fail loudly — the worker has no usable evidence source for
    tailoring and must not silently launch Claude."""
    run = _seed_run(client, tmp_path, monkeypatch)
    input_dir = Path(run["run_dir"]) / "input"

    # Remove the auto-written master_resume.md so the only resume input is
    # the broken DOCX. Also clear the other accepted markdown names just
    # in case (none should be present in a default _seed_run).
    for name in (
        "master_resume.md",
        "resume.md",
        "base_resume.md",
        "original_resume.md",
    ):
        path = input_dir / name
        if path.exists():
            path.unlink()
    (input_dir / "master_resume.docx").write_bytes(b"not a real docx")

    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "failed"
    assert "failed to extract source resume DOCX" in body["error_message"]
    assert "no markdown resume present" in body["error_message"]

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "jobapply: failed to extract source resume DOCX" in log_text


def test_runtime_prompt_references_extracted_markdown_and_mcp():
    """The shipped runtime prompt must reference the extracted markdown
    file the backend writes, and tell Claude to prefer the
    word-document-server MCP for DOCX work."""
    from pathlib import Path

    prompt_path = (
        Path(__file__).resolve().parents[2]
        / "runtime_prompts"
        / "resume_tailoring.md"
    )
    text = prompt_path.read_text(encoding="utf-8")

    assert "input/master_resume_extracted.md" in text
    assert "word-document-server" in text
    # The non-interactive contract must still be present so a future
    # edit does not let the prompt drift back into asking questions.
    assert "Do not ask clarifying questions" in text
    assert "Do not wait for user input" in text
    # The prompt must distinguish DOCX (formatting) from extracted
    # markdown (evidence) so Claude treats them as complementary.
    assert "formatting" in text.lower()
    assert "evidence" in text.lower()


# ---------------------------------------------------------------------------
# Task 079: hardened Word MCP Python env + non-interactive Claude contract
# ---------------------------------------------------------------------------


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[2]


def _word_mcp_doc_text() -> str:
    return (_repo_root() / "docs" / "office_word_mcp_setup.md").read_text(
        encoding="utf-8"
    )


def _install_doc_text() -> str:
    return (_repo_root() / "docs" / "install.md").read_text(encoding="utf-8")


def _runtime_prompt_text() -> str:
    return (
        _repo_root() / "runtime_prompts" / "resume_tailoring.md"
    ).read_text(encoding="utf-8")


def test_word_mcp_docs_recommend_explicit_virtualenv_python():
    """Both docs must walk the reader through Option A (virtualenv) and
    point Claude Code at the in-repo ``.venv`` interpreter rather than a
    PATH-derived ``python``."""
    for text in (_word_mcp_doc_text(), _install_doc_text()):
        lower = text.lower()
        assert "virtualenv" in lower or "venv" in lower
        assert "option a" in lower
        # The virtualenv Python path must appear with its OS-specific
        # interpreter name so a copy/paste user lands on the right
        # binary on either Linux/macOS or Windows.
        assert ".venv/bin/python" in text
        assert r".venv\Scripts\python.exe" in text


def test_word_mcp_docs_recommend_explicit_conda_python():
    """Option B (conda) must appear in both docs and resolve the
    interpreter via ``sys.executable`` so the registered path is the
    real env Python, not whatever ``python`` resolves to next shell."""
    for text in (_word_mcp_doc_text(), _install_doc_text()):
        lower = text.lower()
        assert "option b" in lower
        assert "conda create" in text
        assert "word-mcp" in text
        assert "sys.executable" in text


def test_word_mcp_docs_warn_against_system_python():
    """Docs must explicitly tell the reader not to register the MCP
    with bare ``python`` / ``python3`` on PATH."""
    for text in (_word_mcp_doc_text(), _install_doc_text()):
        # The warning must list both bare interpreter names so neither
        # is implied to be safe.
        lower = text.lower()
        assert "do not rely on" in lower
        # Both names must appear inside the file (covered by the conda
        # / venv setup blocks plus the warning).
        assert "python3" in text
        assert "PATH" in text


def test_word_mcp_docs_avoid_hardcoded_user_paths():
    """No primary command may bake in a developer-specific home path.

    The docs may *describe* such paths in prose (e.g. "paths look like
    /home/<user>/..."), but real shell snippets must use portable
    placeholders or shell variables.
    """
    for path_str, text in (
        ("docs/office_word_mcp_setup.md", _word_mcp_doc_text()),
        ("docs/install.md", _install_doc_text()),
    ):
        # Literal "/home/calvin" / "C:\Users\Calvin" must never appear —
        # those are the most likely hardcoded leak from a local install.
        assert "/home/calvin" not in text, (
            f"{path_str} contains a hardcoded /home/calvin path"
        )
        assert "C:\\Users\\Calvin" not in text, (
            f"{path_str} contains a hardcoded C:\\Users\\Calvin path"
        )
        # Portable variables must be present.
        assert "$WORD_MCP_DIR" in text
        assert "$WORD_MCP_PYTHON" in text
        assert "$WORD_MCP_SERVER" in text


def test_word_mcp_docs_include_linux_macos_verification_commands():
    """Linux/macOS setup must include ``test -x``/``test -f`` checks so
    the reader can confirm the venv Python and server entrypoint are
    actually on disk before registering the MCP."""
    for text in (_word_mcp_doc_text(), _install_doc_text()):
        assert 'test -x "$WORD_MCP_PYTHON"' in text
        assert 'test -f "$WORD_MCP_SERVER"' in text
        # The echo lines surfacing the resolved paths must be present so
        # the user can sanity-check before copying into ``claude mcp add``.
        assert 'echo "$WORD_MCP_PYTHON"' in text
        assert 'echo "$WORD_MCP_SERVER"' in text


def test_word_mcp_docs_include_windows_powershell_verification_commands():
    """Windows PowerShell setup must include ``Test-Path`` checks and
    the matching ``Write-Host`` lines for the same sanity-check pass."""
    for text in (_word_mcp_doc_text(), _install_doc_text()):
        assert "Test-Path $WORD_MCP_PYTHON" in text
        assert "Test-Path $WORD_MCP_SERVER" in text
        assert "Write-Host $WORD_MCP_PYTHON" in text
        assert "Write-Host $WORD_MCP_SERVER" in text


def test_word_mcp_docs_include_claude_mcp_registration_with_placeholders():
    """The ``claude mcp add`` snippet must use the resolved variables so
    a copy/paste reader does not have to guess interpreter paths."""
    for text in (_word_mcp_doc_text(), _install_doc_text()):
        assert "claude mcp add word-document-server" in text
        # The registration block must reference both the explicit
        # interpreter and server-script variables, not bare ``python``.
        assert '"$WORD_MCP_PYTHON"' in text
        assert '"$WORD_MCP_SERVER"' in text
        # Help reference for version-specific flag drift.
        assert "claude mcp add --help" in text


def test_runtime_prompt_blocks_permission_and_clarifying_questions():
    """The runtime prompt must explicitly refuse every shape of
    permission/clarifying prompt a backend run could hang on."""
    text = _runtime_prompt_text()
    for phrase in (
        "non-interactive",
        "Do not ask clarifying questions",
        "Do not wait for user input",
        "Do not ask the user whether to apply changes",
        "Do not ask for permission to edit the resume",
    ):
        assert phrase in text, f"missing phrase: {phrase!r}"


def test_runtime_prompt_grants_run_directory_edit_permission():
    """The prompt must tell Claude the task contract already grants
    permission to create/edit files inside the run directory, so it
    does not stop to ask."""
    text = _runtime_prompt_text()
    assert (
        "task contract already grants permission to create and edit "
        "files inside this run directory"
    ) in text


def test_runtime_prompt_restricts_writes_to_run_directory():
    """The prompt must scope writes to the run directory only — the
    worker enforces this via ``cwd``, but the prompt must reinforce it
    so Claude doesn't drift into project-level edits."""
    text = _runtime_prompt_text()
    assert "Only write inside this run directory" in text


# ---------------------------------------------------------------------------
# Task 092: ATS optimization harness
# ---------------------------------------------------------------------------


def test_runtime_prompt_mentions_ats_optimization():
    """The runtime prompt must call out ATS optimization explicitly."""
    text = _runtime_prompt_text()
    assert "ATS Optimization" in text
    assert "Applicant Tracking System" in text


def test_runtime_prompt_extracts_keyword_classes():
    """The prompt must instruct Claude to classify keywords by importance."""
    text = _runtime_prompt_text()
    assert "required" in text
    assert "preferred" in text
    assert "industry/role-specific" in text
    # The extraction list must cover the named keyword classes.
    for phrase in (
        "exact job title",
        "company name",
        "required skills",
        "preferred skills",
        "tools/technologies",
        "certifications/degrees",
        "domain keywords",
        "repeated phrases",
        "responsibility keywords",
    ):
        assert phrase in text, f"missing phrase: {phrase!r}"


def test_runtime_prompt_forbids_keyword_stuffing():
    """The prompt must explicitly forbid keyword stuffing."""
    text = _runtime_prompt_text()
    assert "Do not keyword-stuff" in text


def test_runtime_prompt_requires_evidence_backed_keywords():
    """ATS keywords must only be inserted when truthful and evidence-backed."""
    text = _runtime_prompt_text()
    assert (
        "Use ATS keywords only when they are truthful and supported by the master"
        in text
    )
    assert (
        "Do not add unsupported skills, certifications, degrees, employers, dates"
        in text
    )


def test_runtime_prompt_requires_standard_section_headings():
    """The prompt must require ATS-safe standard section headings."""
    text = _runtime_prompt_text()
    for heading in (
        "Professional Summary",
        "Skills",
        "Work Experience",
        "Projects",
        "Education",
    ):
        assert heading in text, f"missing section heading: {heading!r}"


def test_runtime_prompt_warns_against_unsafe_formatting():
    """Critical resume content must not live in headers/footers/text boxes/etc."""
    text = _runtime_prompt_text()
    for unsafe in (
        "headers/footers",
        "text boxes",
        "images",
        "graphics",
        "complex tables",
        "multi-column layouts",
    ):
        assert unsafe in text, f"missing formatting warning: {unsafe!r}"


def test_runtime_prompt_requires_acronym_and_full_phrase_usage():
    """The prompt must encourage acronym + full phrase usage when truthful."""
    text = _runtime_prompt_text()
    assert "acronym and full phrase" in text
    assert "Large Language Models (LLMs)" in text
    assert "Applicant Tracking System (ATS)" in text
    assert "Machine Learning (ML)" in text


def test_runtime_prompt_revision_preserves_ats_coverage_and_updates_audit():
    """Revision runs must keep ATS coverage and refresh the ATS audit."""
    text = _runtime_prompt_text()
    assert "preserve ATS-relevant keywords" in text
    assert "update `output/ats_audit.md`" in text


def test_runtime_prompt_requires_ats_audit_output():
    """ats_audit.md must be listed in the required outputs of the prompt."""
    text = _runtime_prompt_text()
    assert "output/ats_audit.md" in text
    # The audit template must include the structured headings the task spec
    # requires so a reader of the prompt can see the contract end-to-end.
    for heading in (
        "## Target Role",
        "## Extracted Keywords",
        "## Keyword Coverage",
        "## Formatting Check",
        "## Risks",
        "## Summary",
    ):
        assert heading in text, f"missing ATS audit heading: {heading!r}"
    # Unsupported-keyword handling must be spelled out.
    assert "Keyword not used because unsupported by evidence" in text


def test_invoke_run_logs_ats_optimization_requested(client, tmp_path, monkeypatch):
    """The worker must log that ATS optimization was requested for the run."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "jobapply: ATS optimization requested" in log_text
    assert "jobapply: ATS audit expected at output/ats_audit.md" in log_text


def test_invoke_run_missing_ats_audit_marks_failed(client, tmp_path, monkeypatch):
    """A fake Claude that writes everything except ats_audit.md must fail."""
    run = _seed_run(client, tmp_path, monkeypatch)
    # Omit ats_audit.md from the write list.
    write_outputs = tuple(name for name in ALL_OUTPUTS if name != "ats_audit.md")
    binary = _write_fake_binary(
        tmp_path,
        exit_code=0,
        write_outputs=write_outputs,
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "failed"
    assert "output/ats_audit.md" in body["error_message"]


def test_invoke_run_with_ats_audit_marks_completed(client, tmp_path, monkeypatch):
    """A fake Claude that writes ats_audit.md plus the other required
    outputs must reach completed."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(
        tmp_path,
        exit_code=0,
        write_outputs=ALL_OUTPUTS,
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"
    assert body["error_message"] is None

    run_dir = Path(body["run_dir"])
    assert (run_dir / "output" / "ats_audit.md").is_file()


def test_invoke_run_dry_run_writes_ats_audit_placeholder(
    client, tmp_path, monkeypatch
):
    """Dry-run mode must populate ats_audit.md so the output contract holds."""
    run = _seed_run(client, tmp_path, monkeypatch)
    monkeypatch.setenv(
        "JOBAPPLY_CLAUDE_BINARY", str(tmp_path / "would_explode_if_used")
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_DRY_RUN", "1")

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    ats_audit = Path(body["run_dir"]) / "output" / "ats_audit.md"
    assert ats_audit.is_file()
    assert "ATS Audit" in ats_audit.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# Task 105: preserve master resume DOCX style
# ---------------------------------------------------------------------------


def _revision_prompt_text() -> Path:
    return (
        _repo_root() / "runtime_prompts" / "resume_revision.md"
    ).read_text(encoding="utf-8")


def _normalize_whitespace(text: str) -> str:
    """Collapse runs of whitespace so prompt-text assertions ignore wrapping.

    Markdown prompts wrap long sentences across newlines for readability;
    these tests care that the *content* is present, not that the exact
    line breaks land in a particular place.
    """
    return " ".join(text.split())


def test_runtime_prompt_preserves_master_docx_style():
    """Tailoring prompt must call out preserving the master DOCX's style."""
    normalized = _normalize_whitespace(_runtime_prompt_text()).lower()
    assert "formatting/style source of truth" in normalized
    assert "preserve the master resume's professional styling" in normalized


def test_runtime_prompt_preserves_heading_colors():
    """Tailoring prompt must specifically call out heading colors."""
    text = _runtime_prompt_text()
    assert "section heading colors" in text
    assert "heading styles/colors" in text


def test_runtime_prompt_preserves_blue_or_simple_colored_headings():
    """Tailoring prompt must explicitly mention blue / simple colored headings."""
    normalized = _normalize_whitespace(_runtime_prompt_text()).lower()
    assert "blue section headers" in normalized
    assert "simple color styling" in normalized


def test_runtime_prompt_copies_or_edits_source_docx():
    """Tailoring prompt must instruct copy/edit of the source DOCX."""
    normalized = _normalize_whitespace(_runtime_prompt_text()).lower()
    assert "copying/editing the source docx" in normalized
    assert "rather than rebuilding a generic resume from scratch" in normalized


def test_runtime_prompt_forbids_plain_text_dump_docx():
    """Tailoring prompt must forbid producing a plain-text dump as the DOCX."""
    text = _runtime_prompt_text()
    assert "Do not create a plain-text dump inside a DOCX" in text


def test_runtime_prompt_balances_style_and_ats():
    """Tailoring prompt must balance style preservation with ATS readability."""
    normalized = _normalize_whitespace(_runtime_prompt_text())
    assert "Preserve visual styling while keeping the resume ATS-readable" in normalized
    # Simple colored headings should be explicitly acceptable.
    assert (
        "Simple colored headings, standard fonts, normal paragraphs, "
        "and bullet lists are acceptable" in normalized
    )


def test_runtime_prompt_warns_against_critical_content_in_unsafe_containers():
    """Tailoring prompt must warn against critical content in headers/text boxes/etc."""
    text = _runtime_prompt_text()
    # The style-preservation section must reiterate the unsafe containers
    # to keep the ATS contract intact even after style preservation lands.
    style_idx = text.lower().find("style preservation vs. ats balance")
    assert style_idx != -1, "style preservation balance section missing"
    style_block = text[style_idx:]
    for unsafe in (
        "headers",
        "footers",
        "text boxes",
        "images",
        "graphics",
        "complex tables",
        "multi-column layouts",
    ):
        assert unsafe in style_block, f"missing unsafe container: {unsafe!r}"


def test_runtime_prompt_word_mcp_style_preservation_instructions():
    """Tailoring prompt must instruct Word MCP usage for style preservation."""
    text = _runtime_prompt_text()
    # The Word MCP usage block must mention preserving paragraph styles,
    # heading styles/colors, and bullet/list styles.
    assert "preserving paragraph styles and run formatting" in text
    assert "preserve heading styles/colors where possible" in text
    assert "preserve bullet/list styles where possible" in text


def test_revision_prompt_preserves_existing_docx_styling():
    """Revision prompt must preserve existing DOCX styling."""
    text = _revision_prompt_text()
    normalized = _normalize_whitespace(text)
    assert (
        "Preserve existing DOCX styling from both the current "
        "tailored draft and the original master resume" in normalized
    )
    assert "Do not restyle the resume unless the user explicitly asks" in normalized
    assert (
        "Apply the requested content changes while preserving existing "
        "DOCX styling" in normalized
    )
    # Specific style elements must appear so future prompt drift cannot
    # silently strip the list.
    for element in (
        "section heading colors",
        "font families",
        "margins and paragraph spacing",
        "bullet indentation",
    ):
        assert element in text, f"missing revision style element: {element!r}"


def test_invoke_run_logs_style_preservation_when_master_docx_present(
    client, tmp_path, monkeypatch
):
    """Worker must log style-preservation lines when a master DOCX exists."""
    run = _seed_run(client, tmp_path, monkeypatch)
    input_dir = Path(run["run_dir"]) / "input"
    _write_master_resume_docx(input_dir)

    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert (
        "jobapply: source DOCX style preservation requested" in log_text
    )
    assert (
        "jobapply: master resume DOCX staged as formatting source" in log_text
    )


def test_invoke_run_skips_style_preservation_log_when_no_master_docx(
    client, tmp_path, monkeypatch
):
    """Without a master DOCX the style-preservation lines must not appear."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    # The style-preservation log is gated on the DOCX being present.
    assert "jobapply: source DOCX style preservation requested" not in log_text
    assert (
        "jobapply: master resume DOCX staged as formatting source" not in log_text
    )


# ---------------------------------------------------------------------------
# Task 107: DOCX template fidelity audit
# ---------------------------------------------------------------------------


def test_runtime_prompt_declares_master_docx_as_template_source_of_truth():
    """The tailoring prompt must name the master DOCX as the template
    source of truth so Claude does not regenerate a generic document
    from scratch."""
    normalized = _normalize_whitespace(_runtime_prompt_text()).lower()
    assert "template source of truth" in normalized


def test_runtime_prompt_says_copy_or_edit_source_docx_workflow():
    """The tailoring prompt must spell out the copy → edit → save
    workflow rather than rebuild-from-scratch."""
    normalized = _normalize_whitespace(_runtime_prompt_text()).lower()
    assert "copy `input/master_resume.docx` as the editable base".lower() in normalized
    assert "replace/tailor text inside the copied document" in normalized
    assert (
        "do not rebuild the resume from scratch unless copying/editing "
        "the source docx fails" in normalized
    )


def test_runtime_prompt_requires_centered_name_header_preservation():
    text = _runtime_prompt_text()
    assert "centered name/header block" in text
    assert "centered contact line / links" in text


def test_runtime_prompt_requires_horizontal_separator_preservation():
    text = _runtime_prompt_text()
    # Both the bulleted preservation list and the workflow language must
    # reference horizontal dividers/separators so future drift cannot
    # silently strip the rule.
    assert "horizontal divider/separator lines" in text
    assert "horizontal separators if present" in text


def test_runtime_prompt_requires_bullet_list_preservation():
    text = _runtime_prompt_text()
    assert "bullet list formatting" in text
    normalized = _normalize_whitespace(text).lower()
    assert (
        "the tailored resume should keep bullet points rather than "
        "converting them to plain paragraphs" in normalized
    )


def test_runtime_prompt_requires_blue_colored_heading_preservation():
    text = _runtime_prompt_text()
    assert "blue or colored section heading style" in text


def test_runtime_prompt_word_mcp_preserves_centered_header_and_separators():
    """The Word MCP usage block must explicitly mention centered header
    alignment and horizontal separators so Claude knows to keep them
    when editing through the MCP tools."""
    text = _runtime_prompt_text()
    assert "preserve centered header alignment" in text
    assert "preserve horizontal separators if present" in text
    assert "replace/tailor content without flattening styles" in text


def test_runtime_prompt_lists_template_fidelity_audit_output():
    """The tailoring prompt must list output/template_fidelity_audit.md
    in the Required Outputs section."""
    text = _runtime_prompt_text()
    assert "output/template_fidelity_audit.md" in text


def test_runtime_prompt_describes_template_fidelity_audit_structure():
    """The prompt must include the audit template so Claude knows the
    required structure (source/output paths, checklist, deviations,
    remediation)."""
    text = _runtime_prompt_text()
    assert "# Template Fidelity Audit" in text
    for section in (
        "## Source Template",
        "## Formatting Preservation Checklist",
        "## Known Deviations",
        "## Remediation",
    ):
        assert section in text, f"missing audit section: {section!r}"
    # The checklist rows must cover each documented preservation feature.
    for row in (
        "Centered name/header block",
        "Centered contact line",
        "Blue/colored section headings",
        "Horizontal divider lines",
        "Bullet lists",
        "Date alignment",
        "Margins",
        "Font family/size consistency",
        "Section spacing",
    ):
        assert row in text, f"missing audit checklist row: {row!r}"


def test_revision_prompt_lists_template_fidelity_audit_output():
    text = _revision_prompt_text()
    assert "output/template_fidelity_audit.md" in text


def test_revision_prompt_preserves_centered_header_and_separators():
    """Revision prompt must mention centered header, colored headings,
    separator lines, and bullet lists in its preservation list."""
    text = _revision_prompt_text()
    for element in (
        "centered name/header block",
        "colored section headings",
        "horizontal separator/divider lines",
        "bullet lists",
        "date alignment",
    ):
        assert element in text, f"missing revision preservation element: {element!r}"


def test_revision_prompt_preserves_layout_and_does_not_restyle():
    """Revision prompt must explicitly preserve DOCX styling and layout
    and refuse to restyle absent a user request."""
    normalized = _normalize_whitespace(_revision_prompt_text())
    assert "Do not restyle the resume unless the user explicitly asks" in normalized
    assert (
        "Apply the requested content changes while preserving existing "
        "DOCX styling and layout" in normalized
    )


def test_word_handoff_prompt_preserves_centered_header_and_separators():
    """The Claude for Word handoff prompt must explicitly preserve the
    centered name/contact header, colored section headings, separator
    lines, and bullet lists."""
    from app.word_handoff import _render_prompt

    text = _render_prompt("Some job description")
    assert "centered name/contact header" in text
    assert "colored section headings" in text
    assert (
        "horizontal separator/divider lines" in text
        or "separator/divider lines" in text
    )
    # Bullet preservation must be called out so Word handoff edits do not
    # flatten lists into paragraphs.
    assert (
        "Keep bullet lists as bullet lists" in text
        or "bullet lists" in text.lower()
    )
    # Centered header preservation must be in the action bullet list, not
    # only in the opening sentence.
    assert "Keep the centered name/contact header block centered" in text


def test_invoke_run_logs_template_fidelity_audit_expected_when_master_docx_present(
    client, tmp_path, monkeypatch
):
    """Worker must log that the template fidelity audit is expected when
    a master DOCX is staged for the run."""
    run = _seed_run(client, tmp_path, monkeypatch)
    input_dir = Path(run["run_dir"]) / "input"
    _write_master_resume_docx(input_dir)

    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert (
        "jobapply: template fidelity audit expected at "
        "output/template_fidelity_audit.md"
    ) in log_text


def test_invoke_run_writes_template_fidelity_audit_deterministically(
    client, tmp_path, monkeypatch
):
    """Task 111: the renderer writes ``template_fidelity_audit.md`` from
    the structured JSON regardless of whether Claude produced one, so
    the "missing audit" warning that task 107 introduced should never
    fire on a successful run."""
    run = _seed_run(client, tmp_path, monkeypatch)
    input_dir = Path(run["run_dir"]) / "input"
    _write_master_resume_docx(input_dir)

    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=ALL_OUTPUTS)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    run_dir = Path(body["run_dir"])
    audit_path = run_dir / "output" / "template_fidelity_audit.md"
    assert audit_path.is_file()
    audit_text = audit_path.read_text(encoding="utf-8")
    assert "Deterministic backend DOCX renderer" in audit_text

    log_text = (run_dir / "run.log").read_text(encoding="utf-8")
    assert "template fidelity audit missing" not in log_text


def test_invoke_run_no_audit_warning_when_no_master_docx(
    client, tmp_path, monkeypatch
):
    """Without a master DOCX there is no template to audit fidelity
    against, so the worker must not log the missing-audit warning."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=ALL_OUTPUTS)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "template fidelity audit missing" not in log_text
    # The "audit expected" line is also gated on the DOCX being present.
    assert (
        "jobapply: template fidelity audit expected at "
        "output/template_fidelity_audit.md"
    ) not in log_text


def test_invoke_run_no_audit_warning_when_claude_writes_audit(
    client, tmp_path, monkeypatch
):
    """When Claude writes the template fidelity audit, the warning must
    not appear."""
    run = _seed_run(client, tmp_path, monkeypatch)
    input_dir = Path(run["run_dir"]) / "input"
    _write_master_resume_docx(input_dir)

    extra_outputs = ALL_OUTPUTS + ("template_fidelity_audit.md",)
    binary = _write_fake_binary(
        tmp_path, exit_code=0, write_outputs=extra_outputs
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    run_dir = Path(body["run_dir"])
    assert (run_dir / "output" / "template_fidelity_audit.md").is_file()
    log_text = (run_dir / "run.log").read_text(encoding="utf-8")
    assert "template fidelity audit missing" not in log_text


# ---------------------------------------------------------------------------
# Task 107: deterministic DOCX style audit helper
# ---------------------------------------------------------------------------


def _build_styled_docx(
    path: Path,
    *,
    centered_header: bool = True,
    centered_contact: bool = True,
    colored_heading: bool = True,
    bullets: bool = True,
) -> Path:
    """Build a DOCX fixture with toggleable visual features.

    The toggles cover the features the audit checklist tracks:
    centered name/header, centered contact line, colored section
    heading, and bullet list paragraphs. Tests opt features in/out so
    the deterministic audit can be exercised against both source-like
    and stripped-output-like documents.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    doc = Document()
    name = doc.add_paragraph()
    name.add_run("Jane Doe").bold = True
    if centered_header:
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = doc.add_paragraph("jane@example.com | linkedin.com/in/jane")
    if centered_contact:
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = doc.add_paragraph("Experience", style="Heading 1")
    if colored_heading:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph("Acme Corp — Staff Engineer", style="Heading 2")
    if bullets:
        doc.add_paragraph("Built distributed systems.", style="List Bullet")
        doc.add_paragraph("Shipped features.", style="List Bullet")
    else:
        doc.add_paragraph("Built distributed systems.")
        doc.add_paragraph("Shipped features.")

    doc.save(str(path))
    return path


def test_docx_style_audit_detects_centered_header(tmp_path):
    from app.docx_style_audit import summarize_docx_style

    path = _build_styled_docx(tmp_path / "centered.docx")
    summary = summarize_docx_style(path)
    assert summary.header_centered is True
    assert summary.contact_centered is True


def test_docx_style_audit_detects_missing_centered_header(tmp_path):
    from app.docx_style_audit import summarize_docx_style

    path = _build_styled_docx(
        tmp_path / "left.docx",
        centered_header=False,
        centered_contact=False,
    )
    summary = summarize_docx_style(path)
    assert summary.header_centered is False
    assert summary.contact_centered is False


def test_docx_style_audit_detects_colored_heading(tmp_path):
    from app.docx_style_audit import summarize_docx_style

    path = _build_styled_docx(tmp_path / "colored.docx")
    summary = summarize_docx_style(path)
    assert summary.has_colored_headings is True
    assert summary.colored_heading_count >= 1


def test_docx_style_audit_detects_missing_colored_heading(tmp_path):
    from app.docx_style_audit import summarize_docx_style

    path = _build_styled_docx(tmp_path / "nocolor.docx", colored_heading=False)
    summary = summarize_docx_style(path)
    assert summary.has_colored_headings is False


def test_docx_style_audit_detects_bullets(tmp_path):
    from app.docx_style_audit import summarize_docx_style

    path = _build_styled_docx(tmp_path / "bullets.docx")
    summary = summarize_docx_style(path)
    assert summary.has_bullets is True
    assert summary.bullet_paragraph_count >= 2


def test_docx_style_audit_detects_missing_bullets(tmp_path):
    from app.docx_style_audit import summarize_docx_style

    path = _build_styled_docx(tmp_path / "nobullets.docx", bullets=False)
    summary = summarize_docx_style(path)
    assert summary.has_bullets is False


def test_docx_style_audit_compare_flags_missing_centered_header(tmp_path):
    """The comparator must surface the regression when the source
    centered the header but the output did not."""
    from app.docx_style_audit import (
        compare_template_fidelity,
        summarize_docx_style,
    )

    source = summarize_docx_style(_build_styled_docx(tmp_path / "src.docx"))
    output = summarize_docx_style(
        _build_styled_docx(
            tmp_path / "out.docx",
            centered_header=False,
            centered_contact=False,
        )
    )
    issues = compare_template_fidelity(source, output)
    features = {issue.feature for issue in issues}
    assert "centered header" in features
    assert "centered contact" in features


def test_docx_style_audit_compare_flags_missing_bullets(tmp_path):
    from app.docx_style_audit import (
        compare_template_fidelity,
        summarize_docx_style,
    )

    source = summarize_docx_style(_build_styled_docx(tmp_path / "src.docx"))
    output = summarize_docx_style(
        _build_styled_docx(tmp_path / "out.docx", bullets=False)
    )
    issues = compare_template_fidelity(source, output)
    features = {issue.feature for issue in issues}
    assert "bullet lists" in features


def test_docx_style_audit_compare_no_issues_when_output_matches(tmp_path):
    from app.docx_style_audit import (
        compare_template_fidelity,
        summarize_docx_style,
    )

    source = summarize_docx_style(_build_styled_docx(tmp_path / "src.docx"))
    output = summarize_docx_style(_build_styled_docx(tmp_path / "out.docx"))
    issues = compare_template_fidelity(source, output)
    assert issues == []


# ---------------------------------------------------------------------------
# Task 108: recruiter review agent
# ---------------------------------------------------------------------------


def test_runtime_prompt_requests_recruiter_review_output():
    """The tailoring prompt must list ``output/recruiter_review.md`` in
    its Required Outputs section so Claude knows to write it."""
    text = _runtime_prompt_text()
    assert "output/recruiter_review.md" in text


def test_runtime_prompt_recruiter_review_role_persona():
    """The recruiter review section must instruct Claude to act as a
    recruiter / hiring manager for the target company."""
    text = _runtime_prompt_text()
    # The Recruiter Review section must call out the multiple
    # personas the review needs to cover.
    assert "## Recruiter Review" in text
    for phrase in (
        "recruiter doing an initial screen",
        "hiring manager doing a technical screen",
        "ATS/human keyword-alignment reviewer",
        "credibility/evidence reviewer",
        "readability/formatting reviewer",
    ):
        assert phrase in text, f"missing recruiter review persona: {phrase!r}"


def test_runtime_prompt_recruiter_review_scorecard():
    """The recruiter review section must include the scorecard rows."""
    text = _runtime_prompt_text()
    for row in (
        "Role fit",
        "Technical keyword alignment",
        "Evidence strength",
        "Recruiter readability",
        "Hiring manager credibility",
        "Seniority/level fit",
        "Formatting/professionalism",
    ):
        assert row in text, f"missing scorecard row: {row!r}"


def test_runtime_prompt_recruiter_review_first_30_second_impression():
    text = _runtime_prompt_text()
    assert "First 30-Second Impression" in text


def test_runtime_prompt_recruiter_review_strengths_and_weaknesses():
    text = _runtime_prompt_text()
    assert "## Strengths" in text
    assert "## Weaknesses / Risks" in text


def test_runtime_prompt_recruiter_review_missing_requirements():
    text = _runtime_prompt_text()
    assert "Missing or Under-emphasized Requirements" in text


def test_runtime_prompt_recruiter_review_suggested_rewrites():
    """The recruiter review section must require suggested rewrites for
    weak lines/bullets so a later revision flow can apply them."""
    text = _runtime_prompt_text()
    assert "Lines or Bullets to Improve" in text
    assert "Suggested rewrite" in text


def test_runtime_prompt_recruiter_review_does_not_invent_company_facts():
    """The recruiter review section must explicitly forbid inventing
    company facts beyond the job description."""
    text = _runtime_prompt_text()
    assert (
        "Do not invent facts about the company beyond what the\njob description states."
        in text
        or "Do not invent facts about the company beyond what the job description states."
        in _normalize_whitespace(text)
    )


def test_runtime_prompt_recruiter_review_includes_overall_recommendation():
    text = _runtime_prompt_text()
    for option in (
        "Strong submit",
        "Submit after minor edits",
        "Needs revision before submit",
        "Do not submit yet",
    ):
        assert option in text, f"missing recommendation option: {option!r}"


def test_revision_prompt_requests_recruiter_review_output():
    """Revision runs must also update output/recruiter_review.md."""
    text = _revision_prompt_text()
    assert "output/recruiter_review.md" in text
    normalized = _normalize_whitespace(text).lower()
    assert (
        "re-review the result as a recruiter/hiring manager and update"
        in normalized
    )


def test_recruiter_review_runtime_prompt_file_exists_and_has_contract():
    """The dedicated recruiter review prompt file must ship with the
    same review contract as the inline section in the tailoring prompt."""
    prompt_path = (
        _repo_root() / "runtime_prompts" / "recruiter_review.md"
    )
    assert prompt_path.is_file(), "recruiter_review.md is not shipped"
    text = prompt_path.read_text(encoding="utf-8")
    # Non-interactive contract
    assert "non-interactive backend job" in text.lower()
    assert "Do not ask clarifying questions" in text
    # Personas
    for phrase in (
        "recruiter doing an initial screen",
        "hiring manager doing a technical screen",
    ):
        assert phrase in text, f"missing persona: {phrase!r}"
    # Scorecard, recommendation, suggested rewrites, do-not-invent
    assert "Scorecard" in text
    assert "Overall Recommendation" in text
    assert "First 30-Second Impression" in text
    assert "Lines or Bullets to Improve" in text
    assert "Suggested rewrite" in text
    assert (
        "Do not invent facts about the company beyond what the\njob description states."
        in text
        or "Do not invent facts about the company beyond what the job description states."
        in _normalize_whitespace(text)
    )


def test_invoke_run_logs_recruiter_review_requested(
    client, tmp_path, monkeypatch
):
    """The worker must log that recruiter review was requested for the run."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert "jobapply: recruiter review requested" in log_text
    assert (
        "jobapply: recruiter review expected at output/recruiter_review.md"
        in log_text
    )


def test_invoke_run_warns_when_recruiter_review_missing(
    client, tmp_path, monkeypatch
):
    """When Claude omits recruiter_review.md, the worker must log a
    warning so the operator can spot the regression. The run still
    completes because the file is requested rather than strictly
    required by the worker today."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=ALL_OUTPUTS)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert (
        "jobapply: warning: recruiter review missing at "
        "output/recruiter_review.md"
    ) in log_text


def test_invoke_run_no_recruiter_review_warning_when_present(
    client, tmp_path, monkeypatch
):
    """When Claude writes recruiter_review.md, the warning must not appear."""
    run = _seed_run(client, tmp_path, monkeypatch)
    extra_outputs = ALL_OUTPUTS + ("recruiter_review.md",)
    binary = _write_fake_binary(
        tmp_path, exit_code=0, write_outputs=extra_outputs
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    run_dir = Path(body["run_dir"])
    assert (run_dir / "output" / "recruiter_review.md").is_file()
    log_text = (run_dir / "run.log").read_text(encoding="utf-8")
    assert "recruiter review missing" not in log_text


def test_invoke_run_dry_run_writes_recruiter_review_placeholder(
    client, tmp_path, monkeypatch
):
    """Dry-run must populate recruiter_review.md so the artifact set is uniform."""
    run = _seed_run(client, tmp_path, monkeypatch)
    monkeypatch.setenv(
        "JOBAPPLY_CLAUDE_BINARY", str(tmp_path / "would_explode_if_used")
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_DRY_RUN", "1")

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    review = Path(body["run_dir"]) / "output" / "recruiter_review.md"
    assert review.is_file()
    assert "Recruiter Review" in review.read_text(encoding="utf-8")


def test_get_run_recruiter_review_returns_content_when_present(
    client, tmp_path, monkeypatch
):
    """``GET /runs/{id}/recruiter-review`` must return the file content."""
    run = _seed_run(client, tmp_path, monkeypatch)
    extra_outputs = ALL_OUTPUTS + ("recruiter_review.md",)
    binary = _write_fake_binary(
        tmp_path, exit_code=0, write_outputs=extra_outputs
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    invoke_resp = client.post(f"/runs/{run['id']}/invoke")
    assert invoke_resp.status_code == 200, invoke_resp.text
    review_resp = client.get(f"/runs/{run['id']}/recruiter-review")
    assert review_resp.status_code == 200, review_resp.text
    body = review_resp.json()
    assert body["run_id"] == run["id"]
    assert body["available"] is True
    assert body["content"]
    assert body["path"] == "output/recruiter_review.md"


def test_get_run_recruiter_review_returns_available_false_when_missing(
    client, tmp_path, monkeypatch
):
    """When the file is absent the endpoint reports available=False
    rather than returning 404, so the UI can render a hint."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=ALL_OUTPUTS)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    invoke_resp = client.post(f"/runs/{run['id']}/invoke")
    assert invoke_resp.status_code == 200, invoke_resp.text
    review_resp = client.get(f"/runs/{run['id']}/recruiter-review")
    assert review_resp.status_code == 200, review_resp.text
    body = review_resp.json()
    assert body["available"] is False
    assert body["content"] is None
    assert body["path"] is None


def test_get_run_recruiter_review_unknown_run_returns_404(client):
    resp = client.get("/runs/does-not-exist/recruiter-review")
    assert resp.status_code == 404


# ---------------------------------------------------------------------------
# Task 112: require actual structured resume output files
# ---------------------------------------------------------------------------


def test_runtime_prompt_mentions_tailored_resume_json():
    """The runtime prompt must explicitly name output/tailored_resume.json
    as a required output file (not just describe the schema)."""
    text = _runtime_prompt_text()
    assert "output/tailored_resume.json" in text
    # The prompt must label the JSON as required, not optional.
    assert "**required**" in text or "is required" in text


def test_runtime_prompt_requires_actually_writing_files():
    """The prompt must forbid describing files in place of writing them.

    The original failure (run d6df714b) was a Claude that summarized the
    expected files in its response and exited 0 without writing them.
    The prompt must make the "write the files" requirement loud and
    explicit so the same failure mode cannot recur.
    """
    text = _runtime_prompt_text()
    normalized = _normalize_whitespace(text)
    assert "Actually write the files" in normalized
    assert "Do not merely describe what each file should contain" in normalized
    assert "Do not end your response until the files have been written" in normalized
    assert "Use shell/file-writing operations if needed" in normalized


def test_runtime_prompt_final_checklist_lists_all_required_files():
    """The Final Verification Checklist must enumerate every required
    output file by path so the model has a concrete list to verify
    against before ending its response."""
    text = _runtime_prompt_text()
    assert "Final Verification Checklist" in text
    checklist_idx = text.find("Final Verification Checklist")
    checklist = text[checklist_idx:]
    for required in (
        "output/tailored_resume.json",
        "output/tailored_resume.md",
        "output/change_log.md",
        "output/claim_audit.md",
        "output/ats_audit.md",
        "output/template_fidelity_audit.md",
        "output/recruiter_review.md",
    ):
        assert required in checklist, f"missing from checklist: {required!r}"


def test_runtime_prompt_includes_structured_resume_schema():
    """The prompt must carry the structured resume JSON schema so Claude
    has the exact shape it must produce on disk."""
    text = _runtime_prompt_text()
    assert "Structured Resume JSON" in text
    # Schema sentinels — header.name + sections + the type enum entries.
    assert '"header"' in text
    assert '"contact_items"' in text
    assert '"sections"' in text
    assert "`summary`" in text
    assert "`experience`" in text


def test_invoke_run_missing_tailored_resume_json_marks_failed(
    client, tmp_path, monkeypatch
):
    """A fake Claude that exits cleanly but writes every output *except*
    tailored_resume.json must fail with a clear missing-file message
    naming the JSON path.

    Mirrors run d6df714b: Claude exit code 0 with no JSON on disk.
    """
    run = _seed_run(client, tmp_path, monkeypatch)
    write_outputs = tuple(
        name for name in ALL_OUTPUTS if name != "tailored_resume.json"
    )
    binary = _write_fake_binary(
        tmp_path, exit_code=0, write_outputs=write_outputs
    )
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "failed"
    message = body["error_message"]
    assert "expected output file missing" in message
    assert "output/tailored_resume.json" in message

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert (
        "jobapply: structured resume JSON expected at "
        "output/tailored_resume.json"
    ) in log_text
    assert (
        "jobapply: expected output file missing: "
        "output/tailored_resume.json"
    ) in log_text


def test_invoke_run_with_tailored_resume_json_passes_validation(
    client, tmp_path, monkeypatch
):
    """A fake Claude that writes the full output set including a valid
    tailored_resume.json must reach completed, and the worker log must
    show the validation step explicitly so operators can see the
    structured JSON was checked."""
    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=ALL_OUTPUTS)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"
    assert body["error_message"] is None

    run_dir = Path(body["run_dir"])
    json_path = run_dir / "output" / "tailored_resume.json"
    assert json_path.is_file()
    assert json_path.read_text(encoding="utf-8").strip()

    log_text = (run_dir / "run.log").read_text(encoding="utf-8")
    assert (
        "jobapply: structured resume JSON expected at "
        "output/tailored_resume.json"
    ) in log_text
    assert "jobapply: validating structured resume JSON" in log_text
