"""Tests for ``backend/app/resume_docx_renderer.py``.

Covers the JSON schema gating, the deterministic DOCX rendering
contract (centered header, colored section headings, real bullet
paragraphs, experience entries with dates), and the template fidelity
audit produced alongside the DOCX.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest


MINIMAL_VALID_PAYLOAD = {
    "header": {
        "name": "Calvin Pieters",
        "contact_items": [
            "calvinpieters@gmail.com",
            "linkedin.com/in/calvin-pieters",
            "Haifa, Israel",
        ],
        "subtitle": "Australian citizen currently based in Israel.",
    },
    "sections": [
        {
            "type": "summary",
            "heading": "PROFESSIONAL SUMMARY",
            "paragraphs": [
                "Engineer building agentic developer tooling and ML systems.",
            ],
        },
        {
            "type": "skills",
            "heading": "SKILLS",
            "groups": [
                {"label": "Languages", "items": ["Python", "TypeScript"]},
            ],
        },
        {
            "type": "experience",
            "heading": "EXPERIENCE",
            "entries": [
                {
                    "title": "Agentic AI Developer Tooling",
                    "organization": "Personal Project",
                    "dates": "2025 – Present",
                    "bullets": [
                        "Built repo-aware developer tooling.",
                        "Shipped MCP integrations.",
                    ],
                },
            ],
        },
        {
            "type": "education",
            "heading": "EDUCATION",
            "entries": [
                {
                    "institution": "Technion – Israel Institute of Technology",
                    "degree": "PhD Candidate, Chemical Engineering",
                    "dates": "2023 – Present",
                    "location": "Israel",
                },
            ],
        },
    ],
    "metadata": {
        "target_company": "Acme",
        "target_job_title": "ML Engineer",
        "generated_for_ats": True,
    },
}


def _write_json(path: Path, payload) -> None:
    path.write_text(json.dumps(payload), encoding="utf-8")


def _stage_run_dir(tmp_path: Path, payload=MINIMAL_VALID_PAYLOAD) -> Path:
    run_dir = tmp_path / "run"
    (run_dir / "output").mkdir(parents=True)
    _write_json(run_dir / "output" / "tailored_resume.json", payload)
    return run_dir


# ---------------------------------------------------------------------------
# Schema / load gating
# ---------------------------------------------------------------------------


def test_validate_resume_payload_accepts_minimal_payload():
    from app.resume_docx_renderer import validate_resume_payload

    validated = validate_resume_payload(MINIMAL_VALID_PAYLOAD)
    assert validated["header"]["name"] == "Calvin Pieters"
    assert len(validated["sections"]) == 4


def test_validate_resume_payload_rejects_missing_name():
    from app.resume_docx_renderer import RendererError, validate_resume_payload

    bad = {
        "header": {"name": ""},
        "sections": [
            {"type": "summary", "heading": "S", "paragraphs": ["x"]},
        ],
    }
    with pytest.raises(RendererError, match="header.name must not be empty"):
        validate_resume_payload(bad)


def test_validate_resume_payload_rejects_empty_sections():
    from app.resume_docx_renderer import RendererError, validate_resume_payload

    bad = {"header": {"name": "X"}, "sections": []}
    with pytest.raises(RendererError, match="at least one section"):
        validate_resume_payload(bad)


def test_validate_resume_payload_rejects_unknown_section_type():
    from app.resume_docx_renderer import RendererError, validate_resume_payload

    bad = {
        "header": {"name": "X"},
        "sections": [{"type": "novelty", "heading": "X"}],
    }
    with pytest.raises(RendererError, match="not one of"):
        validate_resume_payload(bad)


def test_validate_resume_payload_rejects_non_string_bullets():
    from app.resume_docx_renderer import RendererError, validate_resume_payload

    bad = {
        "header": {"name": "X"},
        "sections": [
            {
                "type": "experience",
                "heading": "E",
                "entries": [
                    {"title": "t", "bullets": [123]},
                ],
            }
        ],
    }
    with pytest.raises(RendererError, match="bullets"):
        validate_resume_payload(bad)


def test_load_resume_json_missing_file_raises_clear_error(tmp_path):
    from app.resume_docx_renderer import RendererError, load_resume_json

    with pytest.raises(RendererError, match="expected output file missing"):
        load_resume_json(tmp_path / "tailored_resume.json")


def test_load_resume_json_invalid_json_raises_clear_error(tmp_path):
    from app.resume_docx_renderer import RendererError, load_resume_json

    path = tmp_path / "tailored_resume.json"
    path.write_text("{ not valid json", encoding="utf-8")
    with pytest.raises(RendererError, match="invalid tailored resume JSON"):
        load_resume_json(path)


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------


def test_render_resume_from_run_creates_docx_and_audit(tmp_path):
    from app.resume_docx_renderer import render_resume_from_run

    run_dir = _stage_run_dir(tmp_path)
    result = render_resume_from_run(run_dir)

    assert result.docx_path.is_file()
    assert result.audit_path.is_file()
    assert result.section_count == 4
    # Two bullets in the single experience entry.
    assert result.bullet_count >= 2


def test_render_resume_from_run_missing_json_raises_renderer_error(tmp_path):
    from app.resume_docx_renderer import RendererError, render_resume_from_run

    run_dir = tmp_path / "run"
    (run_dir / "output").mkdir(parents=True)
    with pytest.raises(RendererError, match="tailored_resume.json"):
        render_resume_from_run(run_dir)


def test_rendered_docx_has_centered_header(tmp_path):
    """The first non-empty paragraph (the name) must be centered."""
    from app.resume_docx_renderer import render_resume_from_run
    from docx import Document

    run_dir = _stage_run_dir(tmp_path)
    result = render_resume_from_run(run_dir)

    doc = Document(str(result.docx_path))
    name_paragraph = next(p for p in doc.paragraphs if p.text.strip())
    assert name_paragraph.text.strip() == "Calvin Pieters"
    alignment = name_paragraph.alignment
    assert getattr(alignment, "name", str(alignment)).upper() == "CENTER"


def test_rendered_docx_has_centered_contact_line(tmp_path):
    from app.resume_docx_renderer import render_resume_from_run
    from docx import Document

    run_dir = _stage_run_dir(tmp_path)
    result = render_resume_from_run(run_dir)

    doc = Document(str(result.docx_path))
    nonempty = [p for p in doc.paragraphs if p.text.strip()]
    # name (centered) → contact line (centered) → subtitle (centered) → headings
    contact = nonempty[1]
    assert "calvinpieters@gmail.com" in contact.text
    assert getattr(contact.alignment, "name", "").upper() == "CENTER"


def test_rendered_docx_has_colored_section_headings(tmp_path):
    """Section headings should carry an explicit RGB color (the deterministic
    blue) on at least one run."""
    from app.docx_style_audit import summarize_docx_style
    from app.resume_docx_renderer import render_resume_from_run

    run_dir = _stage_run_dir(tmp_path)
    result = render_resume_from_run(run_dir)
    summary = summarize_docx_style(result.docx_path)
    assert summary.has_colored_headings is True
    assert summary.colored_heading_count >= 1


def test_rendered_docx_uses_real_bullet_paragraphs(tmp_path):
    from app.docx_style_audit import summarize_docx_style
    from app.resume_docx_renderer import render_resume_from_run

    run_dir = _stage_run_dir(tmp_path)
    result = render_resume_from_run(run_dir)
    summary = summarize_docx_style(result.docx_path)
    assert summary.has_bullets is True
    assert summary.bullet_paragraph_count >= 2


def test_rendered_docx_contains_experience_title_and_dates(tmp_path):
    from app.resume_docx_renderer import render_resume_from_run
    from docx import Document

    run_dir = _stage_run_dir(tmp_path)
    result = render_resume_from_run(run_dir)
    doc = Document(str(result.docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Agentic AI Developer Tooling" in text
    assert "2025 – Present" in text
    # Bullets are rendered as separate paragraphs and the text should appear.
    assert "Built repo-aware developer tooling." in text


def test_template_fidelity_audit_calls_out_deterministic_mode(tmp_path):
    from app.resume_docx_renderer import render_resume_from_run

    run_dir = _stage_run_dir(tmp_path)
    result = render_resume_from_run(run_dir, source_docx_relpath="input/master_resume.docx")
    audit_text = result.audit_path.read_text(encoding="utf-8")
    assert "# Template Fidelity Audit" in audit_text
    assert "Deterministic backend DOCX renderer" in audit_text
    assert "Source DOCX: input/master_resume.docx" in audit_text
    # The renderer must record that JSON, not Claude/Word MCP, drove the
    # final DOCX.
    assert "structured JSON" in audit_text
    # The checklist that operators rely on (centered header, colored
    # headings, bullet lists, etc.) must appear.
    for row in (
        "Centered name/header block",
        "Centered contact line",
        "Blue/colored section headings",
        "Horizontal divider lines",
        "Bullet lists",
        "Margins",
    ):
        assert row in audit_text


def test_template_fidelity_audit_handles_missing_source_docx(tmp_path):
    """When the run has no master DOCX, the audit must still be written
    and explicitly note the absence of a source template."""
    from app.resume_docx_renderer import render_resume_from_run

    run_dir = _stage_run_dir(tmp_path)
    result = render_resume_from_run(run_dir)
    audit_text = result.audit_path.read_text(encoding="utf-8")
    assert "code-defined template" in audit_text


# ---------------------------------------------------------------------------
# Worker integration smoke
# ---------------------------------------------------------------------------


def test_invoke_run_logs_deterministic_render_pipeline(
    client, tmp_path, monkeypatch
):
    """Task 111 contract: the worker logs the structured-JSON expectation
    and the deterministic render lines so an operator can see the
    pipeline transitions in ``run.log``."""
    from .test_claude_worker import (
        ALL_OUTPUTS,
        _seed_run,
        _write_fake_binary,
    )

    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=ALL_OUTPUTS)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "completed"

    log_text = (Path(body["run_dir"]) / "run.log").read_text(encoding="utf-8")
    assert (
        "jobapply: structured resume JSON expected at "
        "output/tailored_resume.json" in log_text
    )
    assert (
        "jobapply: rendering DOCX deterministically from structured resume JSON"
        in log_text
    )
    assert "jobapply: rendered output/tailored_resume.docx" in log_text


def test_invoke_run_invalid_json_fails_with_clear_error(
    client, tmp_path, monkeypatch
):
    """A run that produces ``tailored_resume.json`` but with invalid
    structure must fail loudly with a ``invalid tailored resume JSON``
    error so the operator can fix it."""
    from .test_claude_worker import (
        ALL_OUTPUTS,
        _seed_run,
        _write_fake_binary,
    )

    run = _seed_run(client, tmp_path, monkeypatch)
    binary = _write_fake_binary(tmp_path, exit_code=0, write_outputs=ALL_OUTPUTS)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(binary))

    resp = client.post(f"/runs/{run['id']}/invoke")
    assert resp.status_code == 200, resp.text

    # Overwrite the just-rendered JSON with an invalid payload and rerun
    # — the worker truncates run.log and re-validates inputs on every
    # invocation, so the second invoke exercises the failure path.
    json_path = Path(resp.json()["run_dir"]) / "output" / "tailored_resume.json"
    # The second invoke ALSO runs the fake binary, which overwrites the
    # JSON with the (valid) MINIMAL_VALID_RESUME_JSON. To force the
    # failure we make the binary write an invalid JSON instead.
    bad_binary = tmp_path / "fake_claude_bad_json"
    import sys
    import textwrap

    bad_binary.write_text(
        textwrap.dedent(
            f"""\
            #!{sys.executable}
            import sys
            from pathlib import Path
            cwd = Path.cwd()
            out = cwd / "output"
            out.mkdir(parents=True, exist_ok=True)
            for name in {list(ALL_OUTPUTS)!r}:
                if name == "tailored_resume.json":
                    (out / name).write_text(
                        '{{"header": {{}}, "sections": []}}',
                        encoding="utf-8",
                    )
                else:
                    (out / name).write_bytes(b"content\\n")
            sys.exit(0)
            """
        ),
        encoding="utf-8",
    )
    bad_binary.chmod(0o755)
    monkeypatch.setenv("JOBAPPLY_CLAUDE_BINARY", str(bad_binary))

    resp2 = client.post(f"/runs/{run['id']}/invoke")
    assert resp2.status_code == 200, resp2.text
    body = resp2.json()
    assert body["status"] == "failed"
    assert "invalid tailored resume JSON" in body["error_message"]
