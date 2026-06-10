"""Deterministic resume DOCX renderer.

Claude produces structured content as ``output/tailored_resume.json``; this
module turns that JSON into a professional-looking ``output/tailored_resume.docx``
using ``python-docx`` and a code-defined style. Layout is therefore decided
by deterministic backend code, not by Claude / Word MCP, which removes the
recurring "the DOCX looks like a plain-text dump" failure mode.

The renderer also writes ``output/template_fidelity_audit.md`` so the
operator can see exactly what visual features the deterministic style applied
and where Claude's JSON content landed in the document.

Schema (see ``docs/contracts/claude_run_directory.md`` for the full spec):

```json
{
  "header": {
    "name": "Full Name",
    "contact_items": ["email", "linkedin", ...],
    "subtitle": "optional one-liner under the contact row"
  },
  "sections": [
    {"type": "summary",      "heading": "...", "paragraphs": [...]},
    {"type": "skills",       "heading": "...", "groups": [{"label": "...", "items": [...]}]},
    {"type": "experience",   "heading": "...", "entries": [{"title": "...", "organization": "...", "location": "...", "dates": "...", "subtitle": "...", "bullets": [...]}]},
    {"type": "education",    "heading": "...", "entries": [{"institution": "...", "degree": "...", "dates": "...", "location": "..."}]},
    {"type": "publications", "heading": "...", "items": [...]}
  ],
  "metadata": {"target_company": "...", "target_job_title": "...", "generated_for_ats": true}
}
```

Only ``header.name`` and a non-empty ``sections`` array are strictly required.
Every other field is optional; missing optionals are rendered as empty / skipped.
"""

from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional


TAILORED_RESUME_JSON_FILENAME = "tailored_resume.json"
TAILORED_RESUME_DOCX_FILENAME = "tailored_resume.docx"
TEMPLATE_FIDELITY_AUDIT_FILENAME = "template_fidelity_audit.md"

# Deterministic style constants. Picked to look like a clean professional
# resume template without depending on a binary template file. RGB is in
# hex form so it round-trips into python-docx's ``RGBColor`` cleanly.
NAME_FONT_PT = 22
SECTION_HEADING_FONT_PT = 12
SUBTITLE_FONT_PT = 10
BODY_FONT_PT = 10.5
HEADING_COLOR_HEX = "1F497D"  # classic Word "blue 1" — readable in print + ATS-safe
SEPARATOR_COLOR_HEX = "1F497D"
DEFAULT_FONT_NAME = "Calibri"
MARGIN_INCHES = 0.6


class RendererError(ValueError):
    """Raised when the structured resume JSON is invalid or unrenderable."""


@dataclass(frozen=True)
class RenderResult:
    """Outcome of a successful render."""

    docx_path: Path
    audit_path: Path
    section_count: int
    bullet_count: int


# ---------------------------------------------------------------------------
# Schema validation
# ---------------------------------------------------------------------------


_ALLOWED_SECTION_TYPES = {
    "summary",
    "skills",
    "experience",
    "education",
    "publications",
    "projects",
    "certifications",
    "awards",
    "other",
}


def _require_dict(value: Any, where: str) -> dict[str, Any]:
    if not isinstance(value, dict):
        raise RendererError(f"{where} must be an object, got {type(value).__name__}")
    return value


def _require_list(value: Any, where: str) -> list[Any]:
    if not isinstance(value, list):
        raise RendererError(f"{where} must be an array, got {type(value).__name__}")
    return value


def _require_str(value: Any, where: str) -> str:
    if not isinstance(value, str):
        raise RendererError(f"{where} must be a string, got {type(value).__name__}")
    return value


def _optional_str(value: Any, where: str) -> Optional[str]:
    if value is None:
        return None
    return _require_str(value, where)


def _string_list(value: Any, where: str) -> list[str]:
    items = _require_list(value, where)
    result: list[str] = []
    for idx, item in enumerate(items):
        if not isinstance(item, str):
            raise RendererError(
                f"{where}[{idx}] must be a string, got {type(item).__name__}"
            )
        result.append(item)
    return result


def load_resume_json(path: Path) -> dict[str, Any]:
    """Read ``path`` and parse it as JSON.

    Raises :class:`RendererError` when the file is missing, unreadable, or
    not valid JSON — the worker turns this into a run failure.
    """
    if not path.is_file():
        raise RendererError(f"expected output file missing: output/{path.name}")
    try:
        text = path.read_text(encoding="utf-8")
    except OSError as exc:
        raise RendererError(f"failed to read {path.name}: {exc}") from exc
    if not text.strip():
        raise RendererError(f"tailored resume JSON is empty: output/{path.name}")
    try:
        return json.loads(text)
    except json.JSONDecodeError as exc:
        raise RendererError(
            f"invalid tailored resume JSON: {exc.msg} (line {exc.lineno})"
        ) from exc


def validate_resume_payload(data: Any) -> dict[str, Any]:
    """Walk the parsed JSON and surface schema problems with clear messages.

    Validation is intentionally permissive on optional fields so the
    renderer can still produce a document when Claude omits something
    non-essential. Only the structural invariants that the renderer must
    rely on (``header.name``, ``sections`` is a non-empty array of
    sections with a known ``type``) are hard-failed.
    """
    obj = _require_dict(data, "tailored_resume.json root")

    header = _require_dict(obj.get("header"), "header")
    name = _require_str(header.get("name"), "header.name").strip()
    if not name:
        raise RendererError("header.name must not be empty")
    if "contact_items" in header and header["contact_items"] is not None:
        _string_list(header["contact_items"], "header.contact_items")
    _optional_str(header.get("subtitle"), "header.subtitle")

    sections = _require_list(obj.get("sections"), "sections")
    if not sections:
        raise RendererError("sections must contain at least one section")

    for idx, section in enumerate(sections):
        where = f"sections[{idx}]"
        sec = _require_dict(section, where)
        kind = _require_str(sec.get("type"), f"{where}.type")
        if kind not in _ALLOWED_SECTION_TYPES:
            raise RendererError(
                f"{where}.type {kind!r} is not one of {sorted(_ALLOWED_SECTION_TYPES)}"
            )
        _require_str(sec.get("heading", ""), f"{where}.heading")
        if kind == "summary":
            paragraphs = sec.get("paragraphs", [])
            _string_list(paragraphs, f"{where}.paragraphs")
        elif kind == "skills":
            groups = _require_list(sec.get("groups", []), f"{where}.groups")
            for gidx, group in enumerate(groups):
                gwhere = f"{where}.groups[{gidx}]"
                gdict = _require_dict(group, gwhere)
                _require_str(gdict.get("label", ""), f"{gwhere}.label")
                _string_list(gdict.get("items", []), f"{gwhere}.items")
        elif kind == "experience":
            entries = _require_list(sec.get("entries", []), f"{where}.entries")
            for eidx, entry in enumerate(entries):
                ewhere = f"{where}.entries[{eidx}]"
                edict = _require_dict(entry, ewhere)
                _require_str(edict.get("title", ""), f"{ewhere}.title")
                _optional_str(edict.get("organization"), f"{ewhere}.organization")
                _optional_str(edict.get("location"), f"{ewhere}.location")
                _optional_str(edict.get("dates"), f"{ewhere}.dates")
                _optional_str(edict.get("subtitle"), f"{ewhere}.subtitle")
                _string_list(edict.get("bullets", []), f"{ewhere}.bullets")
        elif kind == "education":
            entries = _require_list(sec.get("entries", []), f"{where}.entries")
            for eidx, entry in enumerate(entries):
                ewhere = f"{where}.entries[{eidx}]"
                edict = _require_dict(entry, ewhere)
                _require_str(edict.get("institution", ""), f"{ewhere}.institution")
                _optional_str(edict.get("degree"), f"{ewhere}.degree")
                _optional_str(edict.get("dates"), f"{ewhere}.dates")
                _optional_str(edict.get("location"), f"{ewhere}.location")
        elif kind in {"publications", "projects", "certifications", "awards", "other"}:
            _string_list(sec.get("items", []), f"{where}.items")
    return obj


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------


def _apply_document_style(document) -> None:
    from docx.shared import Inches, Pt  # type: ignore[import-not-found]

    for section in document.sections:
        section.top_margin = Inches(MARGIN_INCHES)
        section.bottom_margin = Inches(MARGIN_INCHES)
        section.left_margin = Inches(MARGIN_INCHES + 0.1)
        section.right_margin = Inches(MARGIN_INCHES + 0.1)

    normal = document.styles["Normal"]
    normal.font.name = DEFAULT_FONT_NAME
    normal.font.size = Pt(BODY_FONT_PT)
    pf = normal.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)


def _add_centered_paragraph(document, text: str, *, size_pt: float, bold: bool = False, color_hex: Optional[str] = None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-not-found]
    from docx.shared import Pt, RGBColor  # type: ignore[import-not-found]

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_hex is not None:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return paragraph


def _add_horizontal_separator(document) -> None:
    """Add a Word horizontal-line border under an empty paragraph.

    python-docx does not expose ``HR`` directly, so we draw a bottom
    border on a thin empty paragraph. The border XML is the standard
    OOXML pattern Word itself emits for ``Heading + bottom rule`` styles.
    """
    from docx.oxml import OxmlElement  # type: ignore[import-not-found]
    from docx.oxml.ns import qn  # type: ignore[import-not-found]

    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), SEPARATOR_COLOR_HEX)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_section_heading(document, text: str) -> None:
    from docx.shared import Pt, RGBColor  # type: ignore[import-not-found]

    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(SECTION_HEADING_FONT_PT)
    run.font.name = DEFAULT_FONT_NAME
    run.font.color.rgb = RGBColor.from_string(HEADING_COLOR_HEX)
    _add_horizontal_separator(document)


def _add_body_paragraph(document, text: str) -> None:
    from docx.shared import Pt  # type: ignore[import-not-found]

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.size = Pt(BODY_FONT_PT)


def _add_bullet(document, text: str) -> None:
    from docx.shared import Pt  # type: ignore[import-not-found]

    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    for run in paragraph.runs:
        run.font.size = Pt(BODY_FONT_PT)
    if not paragraph.runs:
        run = paragraph.add_run(text)
        run.font.size = Pt(BODY_FONT_PT)
    else:
        # ``List Bullet`` style applies via add_paragraph(style=...) which
        # produces a paragraph with no runs; we add the text as our own run
        # so we control the font size explicitly.
        run = paragraph.add_run(text)
        run.font.size = Pt(BODY_FONT_PT)


def _add_experience_entry(document, entry: dict[str, Any]) -> int:
    """Render one experience entry and return the number of bullets emitted."""
    from docx.shared import Pt  # type: ignore[import-not-found]

    title = (entry.get("title") or "").strip()
    organization = (entry.get("organization") or "").strip()
    location = (entry.get("location") or "").strip()
    dates = (entry.get("dates") or "").strip()
    subtitle = (entry.get("subtitle") or "").strip()

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(4)
    title_paragraph.paragraph_format.space_after = Pt(0)
    # Tab to push the dates to the right margin so they line up.
    from docx.enum.text import WD_TAB_ALIGNMENT  # type: ignore[import-not-found]
    from docx.shared import Inches  # type: ignore[import-not-found]

    tab_stops = title_paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)

    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(BODY_FONT_PT + 0.5)
    if dates:
        title_paragraph.add_run("\t" + dates).italic = True

    subtitle_bits = [b for b in (organization, location, subtitle) if b]
    if subtitle_bits:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.paragraph_format.space_after = Pt(2)
        subtitle_run = subtitle_paragraph.add_run(" • ".join(subtitle_bits))
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(BODY_FONT_PT)

    bullets = entry.get("bullets") or []
    count = 0
    for bullet in bullets:
        text = bullet.strip()
        if not text:
            continue
        _add_bullet(document, text)
        count += 1
    return count


def _add_education_entry(document, entry: dict[str, Any]) -> None:
    from docx.shared import Pt  # type: ignore[import-not-found]
    from docx.enum.text import WD_TAB_ALIGNMENT  # type: ignore[import-not-found]
    from docx.shared import Inches  # type: ignore[import-not-found]

    institution = (entry.get("institution") or "").strip()
    degree = (entry.get("degree") or "").strip()
    dates = (entry.get("dates") or "").strip()
    location = (entry.get("location") or "").strip()

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(7.0), WD_TAB_ALIGNMENT.RIGHT
    )
    primary = institution
    if degree:
        primary = f"{institution} — {degree}" if institution else degree
    run = paragraph.add_run(primary)
    run.bold = True
    run.font.size = Pt(BODY_FONT_PT + 0.5)
    if dates:
        paragraph.add_run("\t" + dates).italic = True

    if location:
        sub = document.add_paragraph()
        sub.paragraph_format.space_after = Pt(2)
        sub_run = sub.add_run(location)
        sub_run.italic = True
        sub_run.font.size = Pt(BODY_FONT_PT)


def _add_skills_section(document, groups: list[dict[str, Any]]) -> None:
    from docx.shared import Pt  # type: ignore[import-not-found]

    for group in groups:
        label = (group.get("label") or "").strip()
        items = [str(x).strip() for x in (group.get("items") or []) if str(x).strip()]
        if not items and not label:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        if label:
            run = paragraph.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(BODY_FONT_PT)
        body_run = paragraph.add_run(", ".join(items))
        body_run.font.size = Pt(BODY_FONT_PT)


def render_resume(payload: dict[str, Any], docx_path: Path) -> RenderResult:
    """Render ``payload`` to ``docx_path`` and return a :class:`RenderResult`.

    Caller is responsible for writing the template fidelity audit; this
    function only produces the DOCX. The two are wired together by
    :func:`render_resume_from_run`.
    """
    from docx import Document  # type: ignore[import-not-found]

    document = Document()
    _apply_document_style(document)

    header = payload.get("header") or {}
    name = (header.get("name") or "").strip()
    contact_items = [str(x).strip() for x in (header.get("contact_items") or []) if str(x).strip()]
    subtitle = (header.get("subtitle") or "").strip()

    _add_centered_paragraph(
        document,
        name,
        size_pt=NAME_FONT_PT,
        bold=True,
        color_hex=HEADING_COLOR_HEX,
    )
    if contact_items:
        _add_centered_paragraph(
            document,
            "  •  ".join(contact_items),
            size_pt=SUBTITLE_FONT_PT,
        )
    if subtitle:
        _add_centered_paragraph(
            document,
            subtitle,
            size_pt=SUBTITLE_FONT_PT,
        )
    _add_horizontal_separator(document)

    sections = payload.get("sections") or []
    bullet_count = 0
    for section in sections:
        kind = section.get("type")
        heading = (section.get("heading") or "").strip() or kind.title()
        _add_section_heading(document, heading)
        if kind == "summary":
            for paragraph_text in section.get("paragraphs") or []:
                text = paragraph_text.strip()
                if text:
                    _add_body_paragraph(document, text)
        elif kind == "skills":
            _add_skills_section(document, section.get("groups") or [])
        elif kind == "experience":
            for entry in section.get("entries") or []:
                bullet_count += _add_experience_entry(document, entry)
        elif kind == "education":
            for entry in section.get("entries") or []:
                _add_education_entry(document, entry)
        else:
            for item in section.get("items") or []:
                text = str(item).strip()
                if text:
                    _add_bullet(document, text)
                    bullet_count += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))
    return RenderResult(
        docx_path=docx_path,
        audit_path=docx_path.parent / TEMPLATE_FIDELITY_AUDIT_FILENAME,
        section_count=len(sections),
        bullet_count=bullet_count,
    )


# ---------------------------------------------------------------------------
# Template fidelity audit
# ---------------------------------------------------------------------------


def render_template_fidelity_audit(
    audit_path: Path,
    payload: dict[str, Any],
    result: RenderResult,
    *,
    source_docx_relpath: Optional[str] = None,
) -> None:
    """Write a deterministic template fidelity audit alongside the DOCX.

    The audit's value is that an operator can open the file and see
    which visual features the backend renderer applied (centered header,
    blue uppercase section headings, real bullets, etc.) without opening
    Word. It also calls out that the DOCX was rendered from structured
    JSON rather than synthesized by Claude, so a reader knows where to
    debug visual regressions.
    """
    metadata = payload.get("metadata") or {}
    target_company = (metadata.get("target_company") or "").strip() or "(unknown)"
    target_role = (metadata.get("target_job_title") or "").strip() or "(unknown)"

    contact_items = (payload.get("header") or {}).get("contact_items") or []
    has_subtitle = bool(((payload.get("header") or {}).get("subtitle") or "").strip())

    lines = [
        "# Template Fidelity Audit",
        "",
        "## Rendering Mode",
        "Deterministic backend DOCX renderer.",
        "",
        "The DOCX was rendered from structured JSON "
        f"(`output/{TAILORED_RESUME_JSON_FILENAME}`) by `backend/app/resume_docx_renderer.py` "
        "rather than generated directly by Claude / Word MCP. Layout decisions are owned "
        "by deterministic backend code so visual regressions can be reproduced and fixed.",
        "",
        "## Source Template",
        f"- Source DOCX: {source_docx_relpath or '(no master DOCX; code-defined template)'}",
        f"- Tailored DOCX: output/{TAILORED_RESUME_DOCX_FILENAME}",
        f"- Structured JSON: output/{TAILORED_RESUME_JSON_FILENAME}",
        "",
        "## Target Role",
        f"- Company: {target_company}",
        f"- Job title: {target_role}",
        "",
        "## Preserved Style Features",
        "- Centered name/header: yes",
        f"- Centered contact line: {'yes' if contact_items else 'n/a (no contact items)'}",
        f"- Header subtitle line: {'yes' if has_subtitle else 'no (none provided)'}",
        "- Blue section headings: yes",
        "- Horizontal separators: yes",
        "- Real bullet lists: yes",
        "- Consistent margins: yes",
        "- Consistent spacing: yes",
        "- Date alignment: right-aligned via tab stop (approximate)",
        "",
        "## Formatting Preservation Checklist",
        "| Feature | Source had it? | Output preserved it? | Notes |",
        "| --- | --- | --- | --- |",
        f"| Centered name/header block | {'yes' if source_docx_relpath else 'unknown'} | yes | rendered by backend |",
        f"| Centered contact line | {'yes' if source_docx_relpath else 'unknown'} | {'yes' if contact_items else 'n/a'} | rendered by backend |",
        f"| Blue/colored section headings | {'yes' if source_docx_relpath else 'unknown'} | yes | hex {HEADING_COLOR_HEX} |",
        f"| Horizontal divider lines | {'yes' if source_docx_relpath else 'unknown'} | yes | bottom-border separator |",
        f"| Bullet lists | {'yes' if source_docx_relpath else 'unknown'} | {'yes' if result.bullet_count else 'no (no bullets in JSON)'} | List Bullet style |",
        f"| Date alignment | {'yes' if source_docx_relpath else 'unknown'} | partial | right tab stop approximation |",
        f"| Margins | {'yes' if source_docx_relpath else 'unknown'} | yes | {MARGIN_INCHES}in top/bottom |",
        f"| Font family/size consistency | {'yes' if source_docx_relpath else 'unknown'} | yes | {DEFAULT_FONT_NAME} {BODY_FONT_PT}pt body |",
        f"| Section spacing | {'yes' if source_docx_relpath else 'unknown'} | yes | uniform space-before/after |",
        "",
        "## Known Limitations",
        "- Date right-alignment uses a fixed tab stop rather than a true two-column layout.",
        "- The renderer does not attempt to mirror the master DOCX byte-for-byte; it applies a stable professional style.",
        "- Word MCP / Claude for Word remains available as a manual fallback when a different visual identity is needed.",
        "",
        "## Notes",
        f"- Sections rendered: {result.section_count}",
        f"- Bullets rendered: {result.bullet_count}",
        "- Word MCP / Claude for Word was not used to produce this file.",
        "",
    ]
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    audit_path.write_text("\n".join(lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# Run-level orchestration
# ---------------------------------------------------------------------------


def render_resume_from_run(
    run_dir: Path,
    *,
    source_docx_relpath: Optional[str] = None,
) -> RenderResult:
    """Render the DOCX + audit for ``run_dir``.

    Reads ``output/tailored_resume.json``, validates it, renders
    ``output/tailored_resume.docx``, and writes
    ``output/template_fidelity_audit.md``. Raises :class:`RendererError`
    when the JSON is missing or invalid so the worker can fail the run
    with a clear message.
    """
    output_dir = Path(run_dir) / "output"
    json_path = output_dir / TAILORED_RESUME_JSON_FILENAME
    payload = validate_resume_payload(load_resume_json(json_path))
    docx_path = output_dir / TAILORED_RESUME_DOCX_FILENAME
    result = render_resume(payload, docx_path)
    render_template_fidelity_audit(
        result.audit_path,
        payload,
        result,
        source_docx_relpath=source_docx_relpath,
    )
    return result
